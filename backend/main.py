import os
import io
import json
import re
import asyncio
import time
import random
from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional
from pydantic import BaseModel
from google import genai
from google.genai import types
from PyPDF2 import PdfReader
from docx import Document
import struct
import olefile
import requests
from dotenv import load_dotenv
import difflib
import search_store
from phrasebank import process_phrasebank_rewrite

# API Key Manager (multi-key rotation)
try:
    from api_key_manager import get_api_key_manager
    KEY_MANAGER_AVAILABLE = True
except ImportError:
    KEY_MANAGER_AVAILABLE = False
    print("   [i] api_key_manager not found, using single API key mode")

from harvard_guide import HARVARD_GUIDE
from apa_guide import APA_GUIDE
from vancouver_guide import VANCOUVER_GUIDE

# Load env variables
load_dotenv()

app = FastAPI(title="Writing Tools API")

# Setup CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

API_KEY = os.environ.get("GOOGLE_API_KEY")

def get_client(model: str = None):
    """Get a genai Client using the key manager (with rotation) or fallback to single env var."""
    api_key = None
    if KEY_MANAGER_AVAILABLE:
        manager = get_api_key_manager()
        api_key = manager.get_current_key(model=model)
    if not api_key:
        api_key = API_KEY
    if not api_key:
        raise HTTPException(status_code=500, detail="No API key available. Set GOOGLE_API_KEYS or GOOGLE_API_KEY in .env")
    return genai.Client(api_key=api_key)

MAX_RETRIES = 5
RETRY_BASE_DELAY = 2  # seconds

async def gemini_request_with_retry(client, prompt, model='gemini-3-flash-preview', progress_callback=None, config=None):
    """
    Make a Gemini API request with exponential backoff retry for transient errors.
    Retries on 503 UNAVAILABLE, 429 RESOURCE_EXHAUSTED, and connection errors.
    On 429 errors, rotates to next API key via the key manager.
    Pass `config` (a GenerateContentConfig) to enable features like ThinkingConfig.
    """
    key_manager = get_api_key_manager() if KEY_MANAGER_AVAILABLE else None

    last_error = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            kwargs = dict(model=model, contents=prompt)
            if config is not None:
                kwargs['config'] = config
            response = await asyncio.to_thread(
                client.models.generate_content,
                **kwargs,
            )
            # Track usage AFTER successful call (not before)
            if key_manager:
                new_key = key_manager.increment_usage(model=model)
                if new_key:
                    client = genai.Client(api_key=new_key)
            return response
        except Exception as e:
            error_str = str(e)
            is_rate_limited = any(code in error_str for code in ['429', 'RESOURCE_EXHAUSTED'])
            is_retryable = is_rate_limited or any(code in error_str for code in [
                '503', 'UNAVAILABLE',
                'SSL', 'ConnectionError', 'ConnectionReset',
                'Timeout', 'timeout',
                'ServiceUnavailable',
            ])

            # On rate limit, rotate to next key
            if is_rate_limited and key_manager:
                has_backup = key_manager.mark_exhausted(model=model)
                if has_backup:
                    new_key = key_manager.get_current_key(model=model)
                    if new_key:
                        client = genai.Client(api_key=new_key)
                        if progress_callback:
                            await progress_callback(f'Rate limited — rotated to next API key (attempt {attempt}/{MAX_RETRIES})')

            if is_retryable and attempt < MAX_RETRIES:
                delay = RETRY_BASE_DELAY * (2 ** (attempt - 1)) + random.uniform(0, 1)
                if progress_callback:
                    await progress_callback(f'API temporarily unavailable (attempt {attempt}/{MAX_RETRIES}). Retrying in {delay:.0f}s...')
                await asyncio.sleep(delay)
                last_error = e
            else:
                raise
    raise last_error

def extract_pdf_text(file_bytes: bytes) -> str:
    text = ""
    with io.BytesIO(file_bytes) as f:
        reader = PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_docx_text(file_bytes: bytes) -> str:
    with io.BytesIO(file_bytes) as f:
        doc = Document(f)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

def extract_doc_text(file_bytes: bytes) -> str:
    """Extract text from legacy .doc (Word 97-2003) files using olefile (pure Python)."""
    try:
        f = io.BytesIO(file_bytes)
        ole = olefile.OleFileIO(f)
        
        # The main text stream in .doc files is "WordDocument"
        # But the actual text content is in the "1Table" or "0Table" stream  
        # We'll read the WordDocument stream to get the raw text
        
        if ole.exists('WordDocument'):
            word_stream = ole.openstream('WordDocument').read()
        else:
            raise HTTPException(status_code=400, detail="This file does not appear to be a valid Word .doc file.")
        
        # Read the FIB (File Information Block) to locate text
        # Bytes 24-27 contain flags, bytes 0x01A2 onwards contain text positions
        # For simplicity, try to extract via the compound document text
        
        text_pieces = []
        
        # Method 1: Try to read from the data stream directly
        # The text in a .doc is stored as either ASCII or Unicode
        # We look at ccpText field in FIB at offset 0x004C (76)
        if len(word_stream) > 80:
            ccp_text = struct.unpack_from('<I', word_stream, 0x004C)[0]
            
            # Check if text is Unicode (bit 0 of flags at offset 0x000A)
            flags = struct.unpack_from('<H', word_stream, 0x000A)[0]
            is_complex = not (flags & 0x0004)  # fComplex flag
            
            if not is_complex and ccp_text > 0:
                # Simple file: text starts at offset 0x0200 (512)
                start = 0x0200
                if flags & 0x0100:  # Unicode
                    raw = word_stream[start:start + ccp_text * 2]
                    text_pieces.append(raw.decode('utf-16-le', errors='ignore'))
                else:
                    raw = word_stream[start:start + ccp_text]
                    text_pieces.append(raw.decode('cp1252', errors='ignore'))
        
        # Method 2: If Method 1 got nothing, try brute-force decoding 
        if not text_pieces or not ''.join(text_pieces).strip():
            # Try all text streams
            for stream_name in ['WordDocument']:
                data = ole.openstream(stream_name).read()
                # Skip the FIB header (first 512 bytes) and try to decode
                raw_text = data[512:]
                # Try UTF-16 first, then cp1252
                try:
                    decoded = raw_text.decode('utf-16-le', errors='ignore')
                    # Filter to printable characters
                    cleaned = ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in decoded)
                    if len(cleaned.strip()) > 50:
                        text_pieces = [cleaned]
                except:
                    decoded = raw_text.decode('cp1252', errors='ignore')
                    cleaned = ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in decoded)
                    if len(cleaned.strip()) > 50:
                        text_pieces = [cleaned]

        ole.close()
        
        result = '\n'.join(text_pieces)
        # Clean up control characters but keep newlines
        result = ''.join(c if c.isprintable() or c in '\n\r\t' else '\n' for c in result)
        # Collapse multiple blank lines
        while '\n\n\n' in result:
            result = result.replace('\n\n\n', '\n\n')
        
        return result.strip()
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse .doc file: {str(e)}")

def verify_matches_with_string_search(in_text_citations, references):
    verification_results = {
        "confirmed_matches": [],
        "unmatched_citations": [],
        "unmatched_references": [],
        "duplicate_first_names": {},
        "summary": "Simple first-name match verification completed (case-insensitive)."
    }

    def extract_first_author(citation):
        match = re.search(r'^([A-Za-z\'-]+)', citation.strip('()[] ').split(' et al.')[0].split(' and ')[0].strip())
        if match:
            return match.group(1).strip()
        return None

    def extract_first_ref_author(ref):
        match = re.search(r'^([A-Za-z\s\'-]+)', ref.strip())
        if match:
            author = match.group(1).rstrip(',.').strip()
            return author
        return None

    ref_groups = {}
    for ref in references:
        ref_author = extract_first_ref_author(ref)
        if ref_author:
            ref_groups.setdefault(ref_author.lower(), []).append(ref)

    for first_name, refs in ref_groups.items():
        if len(refs) > 1:
            verification_results["duplicate_first_names"][first_name.capitalize()] = refs

    matched_refs = set()

    for cit in in_text_citations:
        cit_author = extract_first_author(cit)
        if not cit_author:
            verification_results["unmatched_citations"].append(cit)
            continue

        matched = False
        best_ref = None
        for ref in references:
            ref_author = extract_first_ref_author(ref)
            if ref_author and cit_author.lower() == ref_author.lower():
                matched = True
                best_ref = ref
                matched_refs.add(ref)
                break

        if matched:
            verification_results["confirmed_matches"].append({
                "citation": cit,
                "matched_ref": best_ref
            })
        else:
            verification_results["unmatched_citations"].append(cit)

    verification_results["unmatched_references"] = [ref for ref in references if ref not in matched_refs]

    return verification_results

def count_references_and_citations(analysis):
    num_unique_citations = len(analysis.get("in_text_citations", []))
    num_references = len(analysis.get("references", []))
    analysis["num_unique_citations"] = num_unique_citations
    analysis["num_references"] = num_references
    return analysis

def apply_italic_formatting(ref_text: str) -> str:
    """
    Apply <i> tags to parts of a reference that should be italicized
    per academic conventions (Harvard, APA). Uses regex pattern matching.
    
    Rules:
    - Journal/periodical names -> italic
    - Book/report/media/proceedings titles -> italic
    - Article/chapter titles -> NOT italic
    
    Detection priority (most specific first):
    1.  Journal with volume(issue) or bare page range
    1b. Non-standard "In Journal (Vol. X)" format (Google Scholar)
    1c. Conference proceedings: "In Proceedings of..." 
    1d. Newspaper/magazine with section page: "Newspaper, D4."
    1e. Online periodical (no vol/pages, but URL after short source name)
    1f. Advance online publication
    2.  Edited chapter: "In Editor (Ed.), Book Title"
    2b. Media with bracket type: "Title [Film]." or "Title [Video]."
    2c. Dissertation/thesis: "Title [Doctoral dissertation, Univ]"
    3.  Book / report / other (default: italicize title after year)
    """
    import html as html_mod
    ref = html_mod.escape(ref_text)
    
    # Helper: find the period right after the year pattern ")."
    def get_year_dot_pos():
        m = re.search(r'\)\.\s', ref)
        return m.start() + 1 if m else -1
    
    # === 1. JOURNAL / PERIODICAL: detected by volume(issue) OR bare page range ===
    vol_issue = re.search(r',\s*\d{1,4}\s*\(\d{1,4}\)', ref)
    page_range = re.search(r',\s*\d+\s*[-\u2013]\s*\d+', ref)
    periodical_marker = vol_issue or page_range
    if periodical_marker:
        before = ref[:periodical_marker.start()]
        last_dot = before.rfind('. ')
        if last_dot >= 0:
            journal_name = before[last_dot + 2:].rstrip(', ')
            year_dot_pos = get_year_dot_pos()
            if last_dot > year_dot_pos and len(journal_name) > 2 and not journal_name.startswith('http'):
                pos = last_dot + 2
                return ref[:pos] + '<i>' + journal_name + '</i>' + ref[pos + len(journal_name):]
        pass  # Fall through to other checks (e.g. book with date range in title)
    
    # === 1b. NON-STANDARD JOURNAL FORMAT: "In JournalName (Vol. X, No. Y, p. Z)" ===
    in_vol_match = re.search(r'\.\s+In\s+(.+?)\s*\(Vol\.\s*\d+', ref)
    if in_vol_match:
        journal_name = in_vol_match.group(1).rstrip(', ')
        if len(journal_name) > 2:
            start = in_vol_match.start(1)
            end = start + len(in_vol_match.group(1))
            return ref[:start] + '<i>' + journal_name + '</i>' + ref[end:]
    
    # === 1c. CONFERENCE PROCEEDINGS: "In Proceedings of..." ===
    proc_match = re.search(r'\.\s+In\s+(Proceedings\s+of\s+.+?)(?:\s*\(pp?\.|\.|\s*,\s*\d)', ref)
    if proc_match:
        proc_title = proc_match.group(1).rstrip('. ,')
        start = proc_match.start(1)
        end = start + len(proc_match.group(1))
        return ref[:start] + '<i>' + proc_title + '</i>' + ref[end:]
    
    # === 1d. NEWSPAPER/MAGAZINE with section-page: ". Source Name, D4." ===
    section_page = re.search(r',\s*([A-Z]\d{1,3})\s*\.', ref)
    if section_page:
        before = ref[:section_page.start()]
        last_dot = before.rfind('. ')
        if last_dot >= 0:
            source_name = before[last_dot + 2:].rstrip(', ')
            year_dot_pos = get_year_dot_pos()
            if last_dot > year_dot_pos and len(source_name) > 2 and not source_name.startswith('http'):
                pos = last_dot + 2
                return ref[:pos] + '<i>' + source_name + '</i>' + ref[pos + len(source_name):]
    
    # === 2. EDITED CHAPTER: contains "In ... (Ed.), " or "In ... (Eds.), " ===
    ed_match = re.search(r'\bIn\s+.*?\(Eds?\.\)[,.]?\s*', ref)
    if ed_match:
        title_start = ed_match.end()
        rest = ref[title_start:]
        title_end = re.search(r'(?:\.|(?:\s\(pp?\.)|(?:\s\(\d))', rest)
        if title_end:
            title = rest[:title_end.start()]
            if len(title) > 2:
                return ref[:title_start] + '<i>' + title + '</i>' + ref[title_start + len(title):]
        return ref
    
    # === 2b. MEDIA with bracket type: "Title [Film]." or "Title [Video]." etc ===
    # Find the title text immediately before a bracket type descriptor
    media_bracket = re.search(
        r'\[(Film|Video|Motion picture|TV series|TV series episode|'
        r'Webinar|Audio podcast episode|Podcast episode|Song|Album|'
        r'Radio broadcast|Infographic|PowerPoint slides?|Data set|Map|'
        r'Unpublished manuscript|Software|App)\]',
        ref, re.IGNORECASE
    )
    if media_bracket:
        # Title is text between the last "). " (or second-last) and the bracket
        bracket_pos = media_bracket.start()
        before_bracket = ref[:bracket_pos].rstrip()
        # Find where the title starts (after last "). " sequence)
        title_start_match = re.search(r'\)\.\s+', before_bracket)
        if title_start_match:
            # Handle "(Director). (Year). Title" — find the LAST "). "
            last_paren_dot = None
            for m in re.finditer(r'\)\.\s+', before_bracket):
                last_paren_dot = m
            if last_paren_dot:
                title = before_bracket[last_paren_dot.end():].strip()
                if len(title) > 2:
                    start = last_paren_dot.end()
                    return ref[:start] + '<i>' + title + '</i>' + ref[start + len(title):]
    
    # === 2c. DISSERTATION / THESIS ===
    diss_match = re.search(
        r'\)\.\s+(.+?)(?:\s*\(Publication No\.|\s*\[(Doctoral|Master|PhD)\s+(dissertation|thesis))',
        ref, re.IGNORECASE
    )
    if diss_match:
        title = diss_match.group(1).rstrip('. ,')
        if len(title) > 2:
            start = diss_match.start(1)
            end = start + len(diss_match.group(1))
            return ref[:start] + '<i>' + title + '</i>' + ref[end:]
    
    # === 1f. ADVANCE ONLINE PUBLICATION (must be before 1e to take priority) ===
    # Search specifically within after-year text for "Journal Name. Advance online publication"
    year = re.search(r'\((?:\d{4}[a-z]?(?:,\s+\w+(?:\s+\d{1,2})?(?:[\u2013\-]\d{1,2})?)?|n\.d\.)\)\.?\s*', ref)
    if year:
        after_year = ref[year.end():]
        aop_in_after = re.search(r'^.+?\.\s+(.+?)\.\s+Advance\s+online\s+publication', after_year)
        if aop_in_after:
            journal_name = aop_in_after.group(1).strip()
            if len(journal_name) > 2:
                start = year.end() + aop_in_after.start(1)
                end = start + len(aop_in_after.group(1))
                return ref[:start] + '<i>' + journal_name + '</i>' + ref[end:]
    
    # === 1e. ONLINE PERIODICAL (no vol/issue/pages, but URL after short source name) ===
    year = re.search(r'\((?:\d{4}[a-z]?(?:,\s+\w+(?:\s+\d{1,2})?(?:[\u2013\-]\d{1,2})?)?|n\.d\.)\)\.?\s*', ref)
    if year:
        after_year = ref[year.end():]
        online_match = re.search(
            r'^(.+?)\.\s+([A-Z][^.]{2,80}?)\.\s+(https?://|Retrieved\s)',
            after_year
        )
        if online_match:
            source_name = online_match.group(2).strip()
            source_words = source_name.split()
            if 1 <= len(source_words) <= 6:
                start = year.end() + online_match.start(2)
                end = start + len(online_match.group(2))
                return ref[:start] + '<i>' + source_name + '</i>' + ref[end:]
    
    # === 3. BOOK / REPORT / OTHER (default): italicize the title after the year ===
    if year:
        after_year = ref[year.end():]
        segments = re.split(r'(?<=\.)\s+(?=[A-Z])', after_year, maxsplit=3)
        
        if len(segments) >= 2:
            first = segments[0].rstrip('.')
            if len(first) > 2:
                pos = year.end()
                return ref[:pos] + '<i>' + first + '</i>' + ref[pos + len(first):]
        elif len(segments) == 1:
            title = segments[0].rstrip('.')
            if len(title) > 2:
                pos = year.end()
                return ref[:pos] + '<i>' + title + '</i>' + ref[pos + len(title):]
    
    return ref


def extract_verbatim_references(full_text: str, ai_references: list) -> dict:
    """
    For each reference identified by the AI, find the closest verbatim match
    in the original document text. Includes safeguards against DOI bleeding
    and cross-reference mixups.
    
    Returns a dict mapping AI reference -> verbatim source text with confidence.
    """
    from difflib import SequenceMatcher
    
    # ─── STEP 1: Isolate the reference section ───
    # Find where the reference list starts (look for common headings)
    ref_section = full_text
    ref_heading_patterns = [
        r'(?i)\n\s*(references|bibliography|works\s+cited|reference\s+list)\s*\n',
    ]
    ref_start_idx = None
    for pattern in ref_heading_patterns:
        match = re.search(pattern, full_text)
        if match:
            ref_start_idx = match.end()
            break
    
    if ref_start_idx is not None:
        ref_section = full_text[ref_start_idx:]
    
    # ─── STEP 2: Smart atomic splitting ───
    # Split the reference section into individual reference entries.
    # A new reference starts with a line that looks like an author name:
    #   "Surname, I." or "Surname, Initial" or "[1]" or "1."
    # Lines starting with DOI/URL/http are CONTINUATIONS, not new entries.
    
    ref_start_pattern = re.compile(
        r'^(?:'
        r'[A-Z][a-zA-Zà-öø-ÿ\'\-]+\s*,'       # Surname followed by comma (e.g., "Smith,")
        r'|\[\d+\]'                             # Numbered reference [1]
        r'|\d+\.\s+[A-Z]'                       # Numbered reference "1. Author"
        r'|[A-Z][a-zA-Zà-öø-ÿ\'\-]+\s+\('       # Surname followed by "(" (e.g., "Smith (2020)")
        r'|[A-Z]\.?\s*,?\s*\('                   # Single letter + date (e.g., "A. (2020)" or "A (2020)")
        r'|[A-Z]\.?\s*,'                         # Single letter + comma (e.g., "A, B." or "A.,")
        r'|\(\d{4}\)'                             # Year-first format "(2020) Report..."
        r')',
    )

    # Separate pattern for organizational names — requires a year nearby to avoid matching titles
    org_name_pattern = re.compile(r'^[A-Z][a-zA-Zà-öø-ÿ\'\-]+\s+[A-Z]')
    
    # Patterns that indicate a CONTINUATION line (never starts a new reference)
    continuation_pattern = re.compile(
        r'^(?:'
        r'https?://'                             # URL
        r'|[Dd]oi[\s.:]+'                        # DOI (case variants)
        r'|[Aa]vailable\s+at'                    # "Available at:"
        r'|[Aa]ccessed'                          # "Accessed 12 March"
        r'|pp?\.\s*\d'                           # "p. 123" or "pp. 45-67"
        r'|[Vv]ol\.|[Ii]ssue|[Rr]etrieved'      # Volume/Issue/Retrieved
        r'|(?:The|A|An|In|On|Of|For|And|Their|Its|Effects?|Impact)\s'  # Common title-start words
        r'|[a-z]'                                # Starts with lowercase letter (always continuation)
        r'|["\'\u201c\u2018]'                    # Starts with opening quotation mark (title)
        r')'
    )
    
    # Year pattern to validate organizational name lines
    has_year = re.compile(r'\b(?:19|20)\d{2}\b')
    
    lines = ref_section.split('\n')
    atomic_refs = []
    current_ref_lines = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        is_continuation = continuation_pattern.match(stripped)
        
        # For "two capitalized words" (org names), only treat as new ref if a year is found nearby
        is_org_name = org_name_pattern.match(stripped) and not is_continuation
        if is_org_name:
            # Only accept as new ref if line contains a year in first 80 chars
            is_org_name = bool(has_year.search(stripped[:80]))
        
        is_new_ref = (ref_start_pattern.match(stripped) or is_org_name) and not is_continuation
        
        if is_new_ref and current_ref_lines:
            # Save the previous complete reference
            atomic_refs.append(' '.join(current_ref_lines))
            current_ref_lines = [stripped]
        else:
            # Continuation of current reference (including DOIs, URLs, etc.)
            current_ref_lines.append(stripped)
    
    # Don't forget the last reference
    if current_ref_lines:
        atomic_refs.append(' '.join(current_ref_lines))
    
    # Filter out very short entries (likely not actual references)
    atomic_refs = [r for r in atomic_refs if len(r) > 20]

    # ─── POST-SPLIT PASS: detect merged references within a single entry ───
    # Sometimes PDFs put two references on the same line (DOI followed immediately
    # by the next author's name). Detect and split these.
    demerge_pattern = re.compile(
        r'(?<=[0-9])'                          # after a digit (end of DOI/year)
        r'\s+'                                  # whitespace
        r'(?=[A-Z][a-zA-Zà-öø-ÿ\'\-]+\s*,)'   # next author surname + comma
    )
    demerged_refs = []
    for ref in atomic_refs:
        # Check if entry contains multiple DOIs — strong signal of merge
        doi_count = len(re.findall(r'(?:doi[\s.:]+|https?://doi\.org/)', ref, re.IGNORECASE))
        if doi_count >= 2:
            # Split at the boundary between DOI end + next author start
            parts = demerge_pattern.split(ref, maxsplit=1)
            if len(parts) > 1:
                demerged_refs.extend([p.strip() for p in parts if len(p.strip()) > 20])
            else:
                demerged_refs.append(ref)
        else:
            demerged_refs.append(ref)
    atomic_refs = demerged_refs
    
    # ─── STEP 3: Match using author + year compound key ───
    def extract_author_year(text):
        """Extract (first_author_surname, year) from a reference string."""
        author = None
        year = None
        # First author: take text before first comma or parenthesis
        author_match = re.match(r'^[^a-z]*?([A-Z][a-zA-Zà-öø-ÿ\'\-]+)', text.strip())
        if author_match:
            author = author_match.group(1).lower()
        # Year: find a 4-digit year
        year_match = re.search(r'\b(19|20)\d{2}\b', text)
        if year_match:
            year = year_match.group(0)
        return (author, year)
    
    verbatim_map = {}
    used_candidates = {}  # Track which candidate was matched by which AI ref
    
    for ai_ref in ai_references:
        best_score = 0
        best_match = ai_ref  # fallback to AI version if no match found
        ai_author, ai_year = extract_author_year(ai_ref)
        ai_ref_lower = ai_ref.lower().strip()
        
        for candidate in atomic_refs:
            candidate_lower = candidate.lower().strip()
            
            # Quick filter: skip candidates too different in length
            len_ratio = len(candidate_lower) / max(len(ai_ref_lower), 1)
            if len_ratio < 0.3 or len_ratio > 3.0:
                continue
            
            # Compound key check: author AND year must both appear
            cand_author, cand_year = extract_author_year(candidate)
            
            author_ok = (not ai_author) or (not cand_author) or (ai_author == cand_author)
            year_ok = (not ai_year) or (not cand_year) or (ai_year == cand_year)
            
            if not author_ok or not year_ok:
                continue
            
            score = SequenceMatcher(None, ai_ref_lower, candidate_lower).ratio()
            if score > best_score:
                best_score = score
                best_match = candidate
        
        # ─── STEP 4: Conflict detection ───
        conflict = None
        if best_match in used_candidates.values():
            # Another AI reference already matched this same verbatim text
            conflicting_ref = [k for k, v in used_candidates.items() if v == best_match]
            conflict = f"Warning: This verbatim text was also matched by: {conflicting_ref[0][:50]}..."
        
        used_candidates[ai_ref] = best_match
        
        result = {
            "verbatim": best_match,
            "confidence": round(best_score, 2)
        }
        if conflict:
            result["conflict"] = conflict
        
        verbatim_map[ai_ref] = result
    
    return verbatim_map

# ─── PYTHON-GUIDED CITATION EXTRACTION ───

# Reusable pattern components
# Surname: handles Smith, Stokes-Parish, Dall'Ora, O'Brien, Khazaee-Pool, etc.
_SNAME = r"[A-Z\u00C0-\u00D6][a-z\u00E0-\u00F6]+(?:['\u2019\-\u2013\u2014][A-Z\u00C0-\u00D6]?[a-z\u00E0-\u00F6]+)*"
# Year or n.d. or "no date"
_YEAR = r"(?:\d{4}[a-z]?|n\.d\.|no date)"
# Multi-word corporate author: "NHS England", "Survey Coordination Centre", "Nursing and Midwifery Council"
# Requires 2+ words, each starting with a letter, allowing "and"/"of"/"for" etc. between
_CORP = r"(?:[A-Z\u00C0-\u00D6][A-Za-z\u00C0-\u00F6'\-]*\s+)+(?:(?:and|of|for|the)\s+)*[A-Z\u00C0-\u00D6][A-Za-z\u00C0-\u00F6'\-]*(?:\s+(?:(?:and|of|for|the)\s+)*[A-Z\u00C0-\u00D6][A-Za-z\u00C0-\u00F6'\-]*)*"

# Phase 1: Detect multi-citation blocks (parens with semicolons + years)
MULTI_CITATION_PATTERN = re.compile(
    r'\([^()]*\b\d{4}[a-z]?\b[^()]*;[^()]*\b\d{4}[a-z]?\b[^()]*\)'
)

# Phase 2: Individual citation patterns (most specific first)
CITATION_PATTERNS = [
    # ── Author-Date Parenthetical ──
    # Org abbreviation: (WHO, 2020) or (CDC, 2020) or (WHO, n.d.)
    (re.compile(rf'\(\s*[A-Z]{{2,}}[,\s]+{_YEAR}\s*\)'), 'ORG_ABBREV'),

    # With initials: (J. Smith, 2020)
    (re.compile(rf'\(\s*[A-Z]\.\s+{_SNAME}[,\s]+{_YEAR}\s*\)'), 'INITIALS'),

    # Et al.: (Stokes-Parish et al., 2020) or (Dall'Ora et al., 2019, p. 45)
    (re.compile(rf'\(\s*{_SNAME}\s+et\s+al\.?\s*[,\s]*{_YEAR}(?:[,\s]+pp?\.\s*[\d\-\u2013]+)?\s*\)'), 'PAR_ETAL'),

    # Two authors (and/&): (Smith and Jones, 2020) or (Smith & Jones, 2020)
    (re.compile(rf'\(\s*{_SNAME}\s+(?:and|&)\s+{_SNAME}[,\s]+{_YEAR}(?:[,\s]+pp?\.\s*[\d\-\u2013]+)?\s*\)'), 'PAR_TWO'),

    # Secondary referencing: (Ecott, 2002, cited in Wilson, 2009) or (West et al., 2007, quoted in Birch, 2017, p. 17)
    (re.compile(rf'\(\s*{_SNAME}(?:\s+et\s+al\.?)?\s*[,\s]*{_YEAR}[,\s]+(?:cited|quoted)\s+in\s+{_SNAME}(?:\s+et\s+al\.?)?\s*[,\s]*{_YEAR}(?:[,\s]+pp?\.\s*[\d\-\u2013]+)?\s*\)'), 'PAR_SECONDARY'),

    # Multi-word corporate parenthetical: (NHS England, 2025) or (Survey Coordination Centre, 2025)
    (re.compile(rf'\(\s*{_CORP}[,\s]+{_YEAR}(?:[,\s]+pp?\.\s*[\d\-\u2013]+)?\s*\)'), 'PAR_CORP'),

    # Single author with optional page: (Smith, 2020) or (Smith, 2020, p. 45) or (Smith, n.d.)
    (re.compile(rf'\(\s*{_SNAME}[,\s]+{_YEAR}(?:[,:\s]+(?:pp?\.\s*)?[\d\-\u2013]+)?\s*\)'), 'PAR_SINGLE'),

    # ── Author-Date Narrative ──
    # Et al. narrative: Stokes-Parish et al. (2020)
    (re.compile(rf'{_SNAME}\s+et\s+al\.?\s*\({_YEAR}(?:[,\s]+pp?\.\s*[\d\-\u2013]+)?\)'), 'NAR_ETAL'),

    # Two authors narrative: Smith and Jones (2020)
    (re.compile(rf'{_SNAME}\s+(?:and|&)\s+{_SNAME}\s+\({_YEAR}(?:[,\s]+pp?\.\s*[\d\-\u2013]+)?\)'), 'NAR_TWO'),

    # Multi-word corporate narrative: NHS England (2025) or Nuffield Trust (2025)
    (re.compile(rf'{_CORP}\s+\({_YEAR}(?:[,\s]+pp?\.\s*[\d\-\u2013]+)?\)'), 'NAR_CORP'),

    # Single author narrative: Smith (2020) or Smith (2020, p. 45) or Dall'Ora (n.d.)
    (re.compile(rf'{_SNAME}\s+\({_YEAR}(?:[,\s]+pp?\.\s*[\d\-\u2013]+)?\)'), 'NAR_SINGLE'),

    # ── Numbered Styles (Vancouver/IEEE) ──
    # Mixed/multiple numbers: [1, 3-5, 7]
    (re.compile(r'\[\d+(?:\s*[,\-\u2013]\s*\d+)+\]'), 'NUM_MIXED'),

    # Single number: [1]
    (re.compile(r'\[\d+\]'), 'NUM_SINGLE'),

    # ── MLA Style (Author Page) ──
    # (Smith 45) or (Smith 45-67)
    (re.compile(rf'\(\s*{_SNAME}\s+\d+(?:\s*[\-\u2013]\s*\d+)?\s*\)'), 'MLA_PAGE'),
]

# Patterns to match individual citations inside multi-citation blocks (after semicolon split)
# These don't need parentheses — they match the inner text
INNER_CITATION_PATTERNS = [
    # Org abbreviation: WHO, 2020
    (re.compile(rf'^\s*[A-Z]{{2,}}[,\s]+{_YEAR}\s*$'), 'ORG_ABBREV'),
    # Et al.: Stokes-Parish et al., 2020
    (re.compile(rf'^\s*{_SNAME}\s+et\s+al\.?\s*[,\s]*{_YEAR}'), 'PAR_ETAL'),
    # Two authors: Smith and Jones, 2020 or Smith & Jones, 2020
    (re.compile(rf'^\s*{_SNAME}\s+(?:and|&)\s+{_SNAME}[,\s]+{_YEAR}'), 'PAR_TWO'),
    # Multi-word corporate: NHS England, 2025
    (re.compile(rf'^\s*{_CORP}[,\s]+{_YEAR}'), 'PAR_CORP'),
    # Single author: Smith, 2020
    (re.compile(rf'^\s*{_SNAME}[,\s]+{_YEAR}'), 'PAR_SINGLE'),
    # Just a year (same author continuation): 2020
    (re.compile(r'^\s*\d{4}[a-z]?\s*$'), 'YEAR_ONLY'),
]

def extract_reference_section(full_text: str) -> tuple:
    """
    Split the document into body text and reference section.
    Returns (body_text, reference_text).
    """
    ref_heading_patterns = [
        r'(?i)\n\s*(references|bibliography|works\s+cited|reference\s+list)\s*\n',
    ]
    
    for pattern in ref_heading_patterns:
        match = re.search(pattern, full_text)
        if match:
            body = full_text[:match.start()]
            refs = full_text[match.end():]
            return (body, refs)
    
    # Fallback: if no heading found, return full text as body
    return (full_text, "")

def extract_citations_regex(body_text: str) -> list:
    """
    Extract all in-text citations from the body text using regex patterns.
    Uses two-phase approach for multi-citations:
      Phase 1: detect multi-citation blocks, split by semicolon
      Phase 2: match individual citations against single patterns
    
    Returns list of dicts: [{"text": "(Smith, 2020)", "type": "PAR_SINGLE"}, ...]
    """
    found_citations = []

    # Normalize common Unicode variants from PDF/DOCX extraction
    body_text = body_text.replace('\u00a0', ' ')     # non-breaking space -> space
    body_text = body_text.replace('\u2018', "'")      # left single quote -> apostrophe
    body_text = body_text.replace('\u2019', "'")      # right single quote -> apostrophe
    body_text = body_text.replace('\u201c', '"')      # left double quote -> quote
    body_text = body_text.replace('\u201d', '"')      # right double quote -> quote
    body_text = body_text.replace('\u2013', '-')      # en-dash -> hyphen
    body_text = body_text.replace('\u2014', '-')      # em-dash -> hyphen
    body_text = body_text.replace('\ufb01', 'fi')     # fi ligature
    body_text = body_text.replace('\ufb02', 'fl')     # fl ligature
    seen_texts = set()  # Deduplication
    
    # Track positions already matched to avoid double-matching
    matched_spans = []
    
    def is_overlapping(start, end):
        for s, e in matched_spans:
            if start < e and end > s:
                return True
        return False
    
    # ─── Filter out document cross-references ───
    # Patterns like (Table 1), (Figure 2), (Appendix A)
    # Also handles multiple (Table 1, Figure 2) or (Table 1 and 2) or (Table 1 Figure 2)
    cross_ref_pattern = re.compile(
        r'\(\s*(?:Table|Tab\.|Figure|Fig\.|Appendix|App\.)\s+[A-Za-z0-9]+'
        r'(?:\s*(?:,|and|&)?\s*(?:Table|Tab\.|Figure|Fig\.|Appendix|App\.)?\s*[A-Za-z0-9]+)*\s*\)', 
        re.IGNORECASE
    )
    for match in cross_ref_pattern.finditer(body_text):
        matched_spans.append((match.start(), match.end()))
        
    # Phase 1: Find and split multi-citation blocks
    for match in MULTI_CITATION_PATTERN.finditer(body_text):
        if is_overlapping(match.start(), match.end()):
            continue
        matched_spans.append((match.start(), match.end()))
        
        full_block = match.group(0)
        # Strip outer parentheses and split by semicolon
        inner = full_block[1:-1]  # Remove ( and )
        parts = [p.strip() for p in inner.split(';')]
        
        for part in parts:
            part_clean = part.strip()
            if not part_clean:
                continue
            
            # Try to classify each piece
            classified = False
            for pattern, label in INNER_CITATION_PATTERNS:
                if pattern.search(part_clean):
                    citation_text = f"({part_clean})"
                    if citation_text not in seen_texts:
                        found_citations.append({"text": citation_text, "type": label})
                        seen_texts.add(citation_text)
                    classified = True
                    break
            
            # If no inner pattern matched but it has a year, still include it
            if not classified and re.search(r'\d{4}', part_clean):
                citation_text = f"({part_clean})"
                if citation_text not in seen_texts:
                    found_citations.append({"text": citation_text, "type": "UNKNOWN"})
                    seen_texts.add(citation_text)
    
    # Phase 2: Find standalone citations
    for pattern, label in CITATION_PATTERNS:
        for match in pattern.finditer(body_text):
            if is_overlapping(match.start(), match.end()):
                continue
            
            citation_text = match.group(0).strip()
            if citation_text not in seen_texts:
                found_citations.append({"text": citation_text, "type": label})
                seen_texts.add(citation_text)
                matched_spans.append((match.start(), match.end()))
    
    return found_citations

def cross_validate(python_citations: list, ai_citations: list, ai_references: list) -> dict:
    """
    Compare Python-extracted citations with AI-extracted citations.
    Flags discrepancies — citations found by Python but missed by AI,
    and citations claimed by AI but not found by Python (potential hallucination).
    """
    python_texts = set(c["text"] for c in python_citations)
    ai_texts = set(ai_citations)
    
    # Normalize for comparison (lowercase, strip whitespace)
    python_normalized = {t.lower().strip(): t for t in python_texts}
    ai_normalized = {t.lower().strip(): t for t in ai_texts}
    
    # Find discrepancies
    python_only = []  # Python found but AI missed
    ai_only = []      # AI found but Python didn't (possible hallucination)
    confirmed = []    # Both agree
    
    for norm, original in python_normalized.items():
        if norm in ai_normalized:
            confirmed.append(original)
        else:
            # Fuzzy check: maybe AI formatted slightly differently
            found_match = False
            for ai_norm, ai_orig in ai_normalized.items():
                # Check if core content overlaps (first author + year)
                py_author = re.search(r'[A-Za-z\'\-]{2,}', norm)
                ai_author = re.search(r'[A-Za-z\'\-]{2,}', ai_norm)
                py_year = re.search(r'\d{4}', norm)
                ai_year = re.search(r'\d{4}', ai_norm)
                
                if (py_author and ai_author and py_year and ai_year and 
                    py_author.group().lower() == ai_author.group().lower() and
                    py_year.group() == ai_year.group()):
                    confirmed.append(original)
                    found_match = True
                    break
            
            if not found_match:
                python_only.append(original)
    
    for norm, original in ai_normalized.items():
        if norm not in python_normalized:
            # Check fuzzy match as above
            found_match = False
            for py_norm in python_normalized:
                py_author = re.search(r'[A-Za-z\'\-]{2,}', py_norm)
                ai_author = re.search(r'[A-Za-z\'\-]{2,}', norm)
                py_year = re.search(r'\d{4}', py_norm)
                ai_year = re.search(r'\d{4}', norm)
                
                if (py_author and ai_author and py_year and ai_year and 
                    py_author.group().lower() == ai_author.group().lower() and
                    py_year.group() == ai_year.group()):
                    found_match = True
                    break
            
            if not found_match:
                ai_only.append(original)
    
    return {
        "confirmed_by_both": confirmed,
        "python_only": python_only,
        "ai_only_potential_hallucination": ai_only,
        "python_total": len(python_citations),
        "ai_total": len(ai_citations),
    }

@app.post("/api/verify")
async def verify_citations(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(('.pdf', '.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, DOCX, or DOC.")

    file_bytes = await file.read()

    async def event_stream():

        # Stage 1: Detect file format
        yield f"data: {json.dumps({'stage': 'parsing', 'message': 'Detecting file format...'})}\n\n"
        await asyncio.sleep(0.1)
        
        magic = file_bytes[:8]
        try:
            if magic[:4] == b'%PDF':
                yield f"data: {json.dumps({'stage': 'parsing', 'message': 'Reading PDF document...'})}\n\n"
                await asyncio.sleep(0.1)
                full_text = extract_pdf_text(file_bytes)
            elif magic[:2] == b'PK':
                yield f"data: {json.dumps({'stage': 'parsing', 'message': 'Reading Word document (.docx)...'})}\n\n"
                await asyncio.sleep(0.1)
                full_text = extract_docx_text(file_bytes)
            elif magic[:4] == b'\xd0\xcf\x11\xe0':
                yield f"data: {json.dumps({'stage': 'parsing', 'message': 'Reading legacy Word document (.doc)...'})}\n\n"
                await asyncio.sleep(0.1)
                full_text = extract_doc_text(file_bytes)
            else:
                yield f"data: {json.dumps({'stage': 'error', 'message': 'Unrecognized file format.'})}\n\n"
                return
        except Exception as e:
            yield f"data: {json.dumps({'stage': 'error', 'message': f'Failed to read file: {str(e)}'})}\n\n"
            return

        if not full_text:
            yield f"data: {json.dumps({'stage': 'error', 'message': 'No text could be extracted from the file.'})}\n\n"
            return

        word_count = len(full_text.split())
        yield f"data: {json.dumps({'stage': 'extracted', 'message': f'Extracted {word_count:,} words from document'})}\n\n"
        await asyncio.sleep(0.1)

        # Stage 2: Python citation extraction (deterministic — no hallucination)
        yield f"data: {json.dumps({'stage': 'scanning', 'message': 'Python regex scanning for in-text citations...'})}\n\n"
        await asyncio.sleep(0.1)

        body_text, ref_section_text = extract_reference_section(full_text)
        python_citations = extract_citations_regex(body_text)

        py_count = len(python_citations)
        type_counts = {}
        for c in python_citations:
            t = c["type"]
            type_counts[t] = type_counts.get(t, 0) + 1
        type_summary = ", ".join(f"{v} {k}" for k, v in type_counts.items())
        yield f"data: {json.dumps({'stage': 'scanning', 'message': f'Found {py_count} citations via regex ({type_summary})'})}\n\n"
        await asyncio.sleep(0.1)

        # Stage 3: AI Semantic Analysis (receives pre-extracted data — verify only)
        yield f"data: {json.dumps({'stage': 'analyzing', 'message': 'Sending to Gemini AI for semantic matching...'})}\n\n"
        await asyncio.sleep(0.1)

        try:
            model_name = 'gemini-3-flash-preview'
            client = get_client(model=model_name)
            key_manager = get_api_key_manager() if KEY_MANAGER_AVAILABLE else None

            citations_list = json.dumps([c["text"] for c in python_citations], indent=2)

            prompt = f"""You are a citation verification assistant. Python has already extracted the following in-text citations from a document using regex. Your PRIMARY source of truth is this pre-extracted list.

PRE-EXTRACTED IN-TEXT CITATIONS (found by Python regex — these are confirmed):
{citations_list}

DOCUMENT TEXT (for reference matching, irregularity detection, AND independent scanning):
{full_text[:1000000]}

Your tasks:
1. Extract and list all references from the reference section of the document (typically at the end, under headings like 'References', 'Bibliography', etc.).
2. Match each pre-extracted citation to its corresponding reference.
3. Identify mismatches:
   - Citations without a matching reference (missing references).
   - References without a corresponding citation (unused references).
   - Irregularities: Date mismatches, author name inconsistencies, formatting issues, duplicates.
4. INDEPENDENTLY scan the document for any in-text citations that Python may have MISSED. Report these separately as "ai_additional_citations". These are NOT confirmed — they are warnings for the user to review.

Output in strict JSON format only:
{{
    "in_text_citations": {citations_list},
    "references": ["ref1", "ref2", ...],
    "ai_additional_citations": ["any citations YOU found that are NOT in the pre-extracted list above"],
    "missing_references_for_citations": ["unmatched_citation1", ...],
    "unused_references": ["extra_ref1", ...],
    "irregularities": [
        {{"type": "date_mismatch", "citation": "Author (2020)", "ref": "Author (2019)", "details": "Year differs"}},
        {{"type": "name_mismatch", "citation": "Smith", "ref": "Smyth", "details": "Spelling error"}}
    ],
    "summary": "Brief overview of issues found."
}}
"""

            yield f"data: {json.dumps({'stage': 'analyzing', 'message': 'Gemini is matching citations to references...'})}\n\n"
            
            last_retry_error = None
            for attempt in range(1, MAX_RETRIES + 1):
                try:
                    response = await asyncio.to_thread(
                        client.models.generate_content,
                        model=model_name,
                        contents=prompt,
                    )
                    # Track usage after successful call
                    if key_manager:
                        key_manager.increment_usage(model=model_name)
                    break
                except Exception as e:
                    error_str = str(e)
                    is_rate_limited = any(code in error_str for code in ['429', 'RESOURCE_EXHAUSTED'])
                    is_retryable = is_rate_limited or any(code in error_str for code in [
                        '503', 'UNAVAILABLE',
                        'SSL', 'ConnectionError', 'ConnectionReset', 'Timeout', 'timeout',
                        'ServiceUnavailable',
                    ])

                    # On rate limit, rotate to next key
                    if is_rate_limited and key_manager:
                        has_backup = key_manager.mark_exhausted(model=model_name)
                        if has_backup:
                            new_key = key_manager.get_current_key(model=model_name)
                            if new_key:
                                client = genai.Client(api_key=new_key)
                                yield f"data: {json.dumps({'stage': 'analyzing', 'message': f'Rate limited — rotated to next API key (attempt {attempt}/{MAX_RETRIES})'})}\n\n"

                    if is_retryable and attempt < MAX_RETRIES:
                        delay = RETRY_BASE_DELAY * (2 ** (attempt - 1)) + random.uniform(0, 1)
                        yield f"data: {json.dumps({'stage': 'analyzing', 'message': f'API temporarily unavailable (attempt {attempt}/{MAX_RETRIES}). Retrying in {delay:.0f}s...'})}\n\n"
                        await asyncio.sleep(delay)
                        last_retry_error = e
                    else:
                        raise
        except Exception as e:
            yield f"data: {json.dumps({'stage': 'error', 'message': f'Gemini API error: {str(e)}'})}\n\n"
            return

        yield f"data: {json.dumps({'stage': 'processing', 'message': 'Parsing AI response...'})}\n\n"
        await asyncio.sleep(0.1)

        # Stage 4: Parse response and cross-validate
        try:
            output_text = response.text.strip()
            if output_text.startswith('```json'):
                output_text = output_text[7:].strip()
            if output_text.endswith('```'):
                output_text = output_text[:-3].strip()
            analysis = json.loads(output_text)

            analysis = count_references_and_citations(analysis)

            # Stage 5: Cross-validate Python vs AI
            yield f"data: {json.dumps({'stage': 'validating', 'message': 'Cross-validating Python extraction vs AI analysis...'})}\n\n"
            await asyncio.sleep(0.1)

            validation = cross_validate(
                python_citations,
                analysis.get("in_text_citations", []),
                analysis.get("references", [])
            )
            analysis["cross_validation"] = validation
            analysis["python_citations"] = [c["text"] for c in python_citations]
            analysis["python_citation_types"] = {c["text"]: c["type"] for c in python_citations}

            num_cit = analysis.get("num_unique_citations", 0)
            num_ref = analysis.get("num_references", 0)
            py_only = len(validation.get("python_only", []))
            ai_only = len(validation.get("ai_only_potential_hallucination", []))
            ai_additional = len(analysis.get("ai_additional_citations", []))
            validate_msg = f"Validated {num_cit} citations, {num_ref} references"
            if py_only:
                validate_msg += f" | {py_only} found by Python only"
            if ai_only:
                validate_msg += f" | {ai_only} AI-only (review needed)"
            if ai_additional:
                validate_msg += f" | {ai_additional} additional found by AI"
            yield f"data: {json.dumps({'stage': 'validating', 'message': validate_msg})}\n\n"
            await asyncio.sleep(0.1)

            # Stage 6: String verification
            verify_msg = f"Running string verification on {num_ref} references..."
            yield f"data: {json.dumps({'stage': 'verifying', 'message': verify_msg})}\n\n"
            await asyncio.sleep(0.1)

            verification = verify_matches_with_string_search(
                analysis.get("in_text_citations", []),
                analysis.get("references", []),
            )
            analysis["string_verification"] = verification

            # Stage 7: Extract verbatim references from source document
            yield f"data: {json.dumps({'stage': 'extracting', 'message': 'Extracting verbatim references from source document...'})}\n\n"
            await asyncio.sleep(0.1)

            verbatim_map = extract_verbatim_references(full_text, analysis.get("references", []))
            # Apply italic formatting to each verbatim reference
            for key in verbatim_map:
                plain = verbatim_map[key].get("verbatim", key)
                verbatim_map[key]["verbatim_html"] = apply_italic_formatting(plain)
            analysis["verbatim_references"] = verbatim_map

            yield f"data: {json.dumps({'stage': 'complete', 'message': 'Analysis complete!', 'data': analysis})}\n\n"

        except json.JSONDecodeError as e:
            yield f"data: {json.dumps({'stage': 'error', 'message': f'Error parsing Gemini response: {str(e)}'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'stage': 'error', 'message': f'Unexpected error: {str(e)}'})}\n\n"

    from starlette.responses import StreamingResponse
    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ─── PDF REFERENCE GENERATOR ───

def perform_crossref_lookup(doi: str, metadata: dict, field_sources: dict, expected_title: str = None) -> bool:
    """Attempt to fill metadata via CrossRef API. Returns True if successful.
    If expected_title is provided, rejects the lookup if the CrossRef title doesn't match."""
    try:
        crossref_url = f"https://api.crossref.org/works/{doi}"
        resp = requests.get(crossref_url, timeout=10, headers={
            'User-Agent': 'WritingTools/1.0 (mailto:support@paradoxlabs.com)'
        })
        if resp.status_code == 200:
            data = resp.json().get('message', {})
            
            if expected_title:
                titles = data.get('title', [])
                crossref_title = titles[0] if titles else ""
                if crossref_title:
                    t1 = "".join(c for c in expected_title.lower() if c.isalnum() or c.isspace()).strip()
                    t2 = "".join(c for c in crossref_title.lower() if c.isalnum() or c.isspace()).strip()
                    ratio = difflib.SequenceMatcher(None, t1, t2).ratio()
                    if ratio < 0.6:
                        print(f"[CrossRef] Title mismatch for {doi}. Expected: '{expected_title}'. Got: '{crossref_title}' (Ratio: {ratio:.2f})")
                        return False
            
            authors = data.get('author', [])
            if authors:
                author_parts = []
                for a in authors:
                    family = a.get('family', '')
                    given = a.get('given', '')
                    if family and given:
                        initials = '. '.join(w[0].upper() for w in given.split() if w) + '.'
                        author_parts.append(f"{family}, {initials}")
                    elif family:
                        author_parts.append(family)
                if author_parts:
                    metadata["authors"] = author_parts
                    field_sources["authors"] = "crossref"
            
            titles = data.get('title', [])
            if titles:
                metadata["title"] = titles[0]
                field_sources["title"] = "crossref"
            
            date_parts = data.get('published-print', data.get('published-online', data.get('created', {})))
            if date_parts and 'date-parts' in date_parts:
                parts = date_parts['date-parts'][0]
                if parts:
                    metadata["year"] = str(parts[0])
                    field_sources["year"] = "crossref"
            
            container = data.get('container-title', [])
            if container:
                metadata["source"] = container[0]
                field_sources["source"] = "crossref"
            
            if data.get('volume'):
                metadata["volume"] = str(data['volume'])
                field_sources["volume"] = "crossref"
            if data.get('issue'):
                metadata["issue"] = str(data['issue'])
                field_sources["issue"] = "crossref"
            if data.get('page'):
                metadata["pages"] = str(data['page'])
                field_sources["pages"] = "crossref"
            if data.get('publisher'):
                metadata["publisher"] = data['publisher']
                field_sources["publisher"] = "crossref"
            
            cr_type = data.get('type', '')
            type_map = {
                'journal-article': 'Journal Article',
                'book': 'Book',
                'book-chapter': 'Book Chapter',
                'proceedings-article': 'Conference Paper',
                'report': 'Report',
                'dissertation': 'Dissertation',
            }
            metadata["type"] = type_map.get(cr_type, 'Other')
            if cr_type in type_map:
                field_sources["type"] = "crossref"
                
            return True
    except Exception as e:
        print(f"[CrossRef] Lookup failed for {doi}: {e}")
    return False

async def extract_pdf_metadata(file_bytes: bytes) -> dict:
    """
    Extract metadata from a PDF file using PDF properties + first-3-page text parsing.
    Enrichment priority: CrossRef DOI > Python regex > Gemini AI fallback.
    Returns a dict with: authors, title, year, source, doi, url, volume, issue, pages, publisher, type, field_sources.
    """
    metadata = {
        "authors": None,
        "title": None,
        "year": None,
        "source": None,
        "doi": None,
        "url": None,
        "volume": None,
        "issue": None,
        "pages": None,
        "publisher": None,
        "type": "Other",
    }
    field_sources = {}  # Track where each field came from
    
    with io.BytesIO(file_bytes) as f:
        reader = PdfReader(f)
        
        # ─── Step 1: PDF built-in metadata ───
        pdf_meta = reader.metadata
        if pdf_meta:
            if pdf_meta.get('/Author') or pdf_meta.get('author'):
                raw_author = pdf_meta.get('/Author') or pdf_meta.get('author') or ''
                if raw_author and not any(skip in raw_author.lower() for skip in ['microsoft', 'adobe', 'scanner', 'latex', 'tex', 'kman', 'acrobat', 'pdf']):
                    metadata["authors"] = raw_author
                    field_sources["authors"] = "pdf_metadata"
            
            if pdf_meta.get('/Title') or pdf_meta.get('title'):
                raw_title = pdf_meta.get('/Title') or pdf_meta.get('title') or ''
                if raw_title and len(raw_title) > 3:
                    metadata["title"] = raw_title
                    field_sources["title"] = "pdf_metadata"
            
            if pdf_meta.get('/CreationDate') or pdf_meta.get('creation_date'):
                raw_date = str(pdf_meta.get('/CreationDate') or pdf_meta.get('creation_date') or '')
                year_match = re.search(r'(19|20)\d{2}', raw_date)
                if year_match:
                    metadata["year"] = year_match.group(0)
                    field_sources["year"] = "pdf_metadata"
        
        # Sanity check: if author == title (case insensitive), both are likely bogus
        if (metadata.get("authors") and metadata.get("title") and 
            str(metadata["authors"]).strip().lower() == str(metadata["title"]).strip().lower() and
            field_sources.get("authors") == "pdf_metadata"):
            metadata["authors"] = None
            metadata["title"] = None
            field_sources.pop("authors", None)
            field_sources.pop("title", None)
        
        # ─── Step 2: First 3 pages text parsing (handles cover pages) ───
        pages_to_scan = min(3, len(reader.pages))
        first_pages_text = ''
        for i in range(pages_to_scan):
            page_text = reader.pages[i].extract_text() or ''
            first_pages_text += page_text + '\n'
        
        if first_pages_text.strip():
            # DOI extraction — scan all 3 pages
            doi_match = re.search(
                r'(?:doi[:\s]*|https?://(?:dx\.)?doi\.org/)(10\.\d{4,}/[a-zA-Z0-9.\-_/:()\[\]]+)',
                first_pages_text, re.IGNORECASE
            )
            if doi_match:
                doi = doi_match.group(1)
                # Strip trailing punctuation
                doi = re.sub(r'[\].;,()]+$', '', doi)
                # Strip trailing common words accidentally appended
                doi = re.sub(r'(?i)(Research|Article|Review|Copyright|Downloaded)\b.*$', '', doi)
                metadata["doi"] = doi
                field_sources["doi"] = "text_parsing"
            
            # Title: find best candidate across all 3 pages
            # Always try, even if PDF metadata set one (it may be inaccurate like software names)
            lines = [l.strip() for l in first_pages_text.split('\n') if l.strip()]
            skip_patterns = re.compile(
                r'^('
                r'vol|volume|issue|page|doi|http|www|©|ISSN|ISBN|'
                r'RESEARCH|ARTICLE|REVIEW|Open Access|Creative Commons|'
                r'Correspondence|Author|Received|Accepted|Published|'
                r'Abstract|Background|Introduction|Methods|Results|'
                r'Keywords|Licensed|Check for|updates|Citation|'
                r'This article|The Author|Springer|Elsevier|Wiley|BMC|'
                r'\d+\s*$|et al\.'
                r')', re.IGNORECASE
            )
            best_title = None
            best_score = 0
            for line in lines:
                if len(line) < 15 or len(line) > 300:
                    continue
                if skip_patterns.match(line):
                    continue
                if re.match(r'^\d+\s*$', line):
                    continue
                # Skip header/footer lines containing citation patterns
                # e.g. "Olsen & Bastholm Foresight and Public Health 2:1 (2025) 1-10"
                if re.search(r'\d+:\d+\s*\(\d{4}\)', line):   # volume:issue (year)
                    continue
                if re.search(r'\(\d{4}\)\s*\d+[-–]\d+', line):  # (year) page-range
                    continue
                if re.search(r'\b\d+[-–]\d+\s*$', line):  # ends with page range
                    continue
                score = min(len(line), 150)
                if line[0].isupper() and ':' in line:
                    score += 20
                if score > best_score:
                    best_score = score
                    best_title = line
            if best_title:
                if not metadata["title"]:
                    # No title at all — use what we found
                    metadata["title"] = best_title
                    field_sources["title"] = "text_parsing"
                elif field_sources.get("title") == "pdf_metadata":
                    # Only override PDF metadata if it looks suspicious
                    current = str(metadata["title"]).strip()
                    is_suspicious = (
                        len(current) < 10 or              # Too short to be a real title
                        ' ' not in current or              # Single word (likely software name)
                        current.isupper() and len(current) < 20  # All-caps short string
                    )
                    if is_suspicious:
                        metadata["title"] = best_title
                        field_sources["title"] = "text_parsing"
            
            # Year from text
            if not metadata["year"]:
                year_match = re.search(r'\b(19|20)\d{2}\b', first_pages_text)
                if year_match:
                    metadata["year"] = year_match.group(0)
                    field_sources["year"] = "text_parsing"

            # Volume/Issue/Pages
            vol_match = re.search(r'(?:vol(?:ume)?\.?\s*|(?<=,\s))(\d+)\s*\((\d+)\)', first_pages_text, re.IGNORECASE)
            if vol_match:
                metadata["volume"] = vol_match.group(1)
                metadata["issue"] = vol_match.group(2)
                field_sources["volume"] = "text_parsing"
                field_sources["issue"] = "text_parsing"
            
            pages_match = re.search(r'(?:pp?\.?\s*)(\d+)\s*[-–]\s*(\d+)', first_pages_text, re.IGNORECASE)
            if pages_match:
                metadata["pages"] = f"{pages_match.group(1)}-{pages_match.group(2)}"
                field_sources["pages"] = "text_parsing"

            # URL extraction (if no DOI)
            if not metadata["doi"]:
                url_match = re.search(r'(https?://\S+)', first_pages_text)
                if url_match:
                    metadata["url"] = url_match.group(1).rstrip('.,;)')
                    field_sources["url"] = "text_parsing"

    # ─── Step 3: CrossRef DOI lookup ───
    crossref_success = False
    if metadata["doi"]:
        crossref_success = perform_crossref_lookup(metadata["doi"], metadata, field_sources)
    
    # Track crossref state so the UI can warn the user if it failed
    metadata["crossref_failed"] = not crossref_success
    
    # ─── Step 4: Gemini AI — verify non-CrossRef fills + fill missing fields ───
    # Run AI when: (a) CrossRef failed or missing, OR (b) any field came from regex/pdf_metadata
    has_unverified_fields = any(v in ("text_parsing", "pdf_metadata") for v in field_sources.values())
    has_missing = not metadata["authors"] or not metadata["title"] or not metadata["year"]
    
    if (not crossref_success or has_unverified_fields or has_missing) and first_pages_text.strip():
        # Build context of what Python already found (for verification)
        python_found = {}
        for key in ["authors", "title", "year", "source", "doi"]:
            if metadata.get(key) and field_sources.get(key) in ("text_parsing", "pdf_metadata"):
                python_found[key] = metadata[key]
        
        verification_context = ""
        if python_found:
            verification_context = f"""
Python regex has already extracted these fields (VERIFY them — confirm or correct):
{json.dumps(python_found, indent=2, ensure_ascii=False)}
"""
        
        try:
            model_name = 'gemini-3-flash-preview'
            client = get_client(model=model_name)
            prompt = f"""You are a metadata extraction tool. Extract ONLY what you can see in the document text below.
DO NOT invent, guess, or hallucinate any information. If a field is not clearly visible, return null.
{verification_context}
Extract these fields from the document:
- authors: List of author names in "Surname, Initials." format (e.g. ["Smith, J.", "Jones, A. B."])
- title: The main title of the paper/document (not section headings, not journal name)
- year: The publication year (not copyright year if different)
- source: Journal name or publisher name if visible
- doi: The Digital Object Identifier (e.g. 10.1234/example)

DOCUMENT TEXT (first 3 pages):
{first_pages_text[:8000]}

Respond in strict JSON only:
{{
    "authors": ["author1", "author2"] or null,
    "title": "title text" or null,
    "year": "2025" or null,
    "source": "journal or publisher name" or null,
    "doi": "10.1234/example" or null
}}"""
            
            response = await gemini_request_with_retry(client, prompt, model=model_name)
            
            ai_text = response.text.strip()
            if ai_text.startswith('```json'):
                ai_text = ai_text[7:].strip()
            if ai_text.endswith('```'):
                ai_text = ai_text[:-3].strip()
            ai_data = json.loads(ai_text)
            print(f"[AI Verification] AI returned: {ai_data}")
            
            # If AI found a new DOI, and CrossRef previously failed, try CrossRef again!
            ai_doi = ai_data.get("doi")
            if ai_doi and ai_doi != metadata.get("doi") and not crossref_success:
                print(f"[AI Verification] Retrying CrossRef with new AI-extracted DOI: {ai_doi}")
                # Pass expected title for verification against hallucination
                expected = ai_data.get("title") or metadata.get("title")
                retry_success = perform_crossref_lookup(ai_doi, metadata, field_sources, expected_title=expected)
                if retry_success:
                    metadata["doi"] = ai_doi
                    field_sources["doi"] = "ai_verified"
                    crossref_success = True
                    metadata["crossref_failed"] = False
                else:
                    print(f"[AI Verification] Rejected AI DOI {ai_doi} due to CrossRef lookup failure or title mismatch.")
            
            # For each field: fill if missing, or override if from unreliable source (AI verification)
            # Only do this if CrossRef didn't already fill it (CrossRef is authoritative)
            for key in ["authors", "title", "year", "source", "doi"]:
                ai_value = ai_data.get(key)
                if not ai_value:
                    continue
                
                current_source = field_sources.get(key)
                
                # If source is "crossref" — never override, CrossRef is authoritative
                if current_source == "crossref":
                    continue
                    
                if not metadata.get(key):
                    # Field was missing — AI fills it
                    metadata[key] = str(ai_value) if key in ("year", "volume", "issue") else ai_value
                    field_sources[key] = "ai"
                elif current_source in ("text_parsing", "pdf_metadata"):
                    # Field was from unreliable source — AI verifies/corrects it
                    if key in ("year", "volume", "issue"):
                        ai_value = str(ai_value)
                    metadata[key] = ai_value
                    field_sources[key] = "ai_verified"
                
        except Exception as e:
            print(f"[AI Verification] FAILED: {type(e).__name__}: {e}")
            metadata["ai_warning"] = f"AI verification failed: {type(e).__name__}. Metadata may be inaccurate."
    
    # ─── Step 5: Classify type if not set by CrossRef ───
    if metadata["type"] == "Other":
        metadata["type"] = classify_source_type(metadata)
    
    # Ensure authors is a list
    if isinstance(metadata["authors"], str):
        raw = metadata["authors"]
        if '; ' in raw:
            metadata["authors"] = [a.strip() for a in raw.split(';') if a.strip()]
        elif ' and ' in raw or ' & ' in raw:
            raw = raw.replace(' & ', ' and ')
            metadata["authors"] = [a.strip() for a in raw.split(' and ') if a.strip()]
        else:
            metadata["authors"] = [raw.strip()]
    
    metadata["field_sources"] = field_sources
    return metadata


def classify_source_type(metadata: dict) -> str:
    """Classify the source type based on available metadata fields."""
    title = (metadata.get("title") or "").lower()
    source = (metadata.get("source") or "").lower()
    
    if metadata.get("volume") or metadata.get("issue"):
        return "Journal Article"
    if metadata.get("doi"):
        return "Journal Article"  # Most DOI content is journal articles
    # Check for dissertation/thesis keywords in title
    if any(kw in title for kw in ("dissertation", "thesis")):
        return "Dissertation"
    # Check for report
    if metadata.get("report_number") or any(kw in title for kw in ("report", "working paper", "technical report")):
        return "Report"
    # Check for newspaper (has day_month and a source name)
    if metadata.get("day_month") and source:
        return "Newspaper Article"
    if metadata.get("publisher"):
        return "Book"
    if metadata.get("url"):
        return "Web Page"
    return "Other"


def condense_pages(pages_str: str) -> str:
    """Condense page ranges for Vancouver style (e.g., 117-119 → 117-9, 301-307 → 301-7).
    Leaves non-numeric pages unchanged (e.g., 34A-37A stays as-is)."""
    match = re.match(r'^(\d+)\s*[-–]\s*(\d+)$', pages_str.strip())
    if not match:
        return pages_str  # Not a simple numeric range, return as-is
    start, end = match.group(1), match.group(2)
    if len(start) != len(end) or len(start) <= 1:
        return f"{start}-{end}"
    # Find where digits start to differ and truncate
    for i in range(len(start)):
        if start[i] != end[i]:
            return f"{start}-{end[i:]}"
    return f"{start}-{end}"  # Identical numbers, keep as-is


def format_reference(metadata: dict, style: str = "harvard") -> dict:
    """
    Format metadata into a reference string in the specified style.
    Returns { "formatted": "plain text", "formatted_html": "with <i> tags", "metadata": {...} }
    """
    authors = metadata.get("authors") or ["Unknown Author"]
    title = metadata.get("title") or "Untitled"
    year = metadata.get("year") or "n.d."
    source = metadata.get("source")
    volume = metadata.get("volume")
    issue = metadata.get("issue")
    pages = metadata.get("pages")
    doi = metadata.get("doi")
    url = metadata.get("url")
    publisher = metadata.get("publisher")
    ref_type = metadata.get("type", "Other")
    
    # Format author string
    if len(authors) == 1:
        author_str = authors[0]
    elif len(authors) == 2:
        if style == "apa":
            author_str = f"{authors[0]} & {authors[1]}"
        else:
            author_str = f"{authors[0]} and {authors[1]}"
    elif len(authors) <= 20:
        if style == "apa":
            author_str = ', '.join(authors[:-1]) + f', & {authors[-1]}'
        else:
            author_str = ', '.join(authors[:-1]) + f' and {authors[-1]}'
    else:
        # APA 7th: list first 19, then ... then last
        author_str = ', '.join(authors[:19]) + f', ... {authors[-1]}'

    # ─── Vancouver author formatting ───
    if style == "vancouver":
        # Vancouver (NLM): Surname Initials (no periods), comma-separated.
        # NLM lists ALL authors by default, ending with a period.
        van_authors = []
        for a in authors:
            # If already in "Surname AB" format, keep as-is
            parts = a.strip().split(',')
            if len(parts) >= 2:
                surname = parts[0].strip()
                initials = parts[1].strip().replace('.', '').replace(' ', '')
                van_authors.append(f"{surname} {initials}")
            else:
                van_authors.append(a.strip())
        author_str = ', '.join(van_authors) + '.'
    
    # ─── Harvard author formatting ───
    author_str_html = author_str  # Default: same as plain text
    if style == "harvard":
        if len(authors) >= 4:
            author_str = f"{authors[0]} et al."
            author_str_html = f"{authors[0]} <i>et al.</i>"
        # Harvard uses "no date" instead of "n.d."
        if year == "n.d.":
            year = "no date"
    
    # Build location string (vol, issue, pages, DOI/URL)
    location_parts = []
    if volume and issue:
        location_parts.append(f"{volume}({issue})")
    elif volume:
        location_parts.append(volume)
    if pages:
        if location_parts and style == "harvard":
            location_parts.append(f", pp. {pages}")  # Harvard always uses pp.
        elif location_parts:
            location_parts.append(f", {pages}")
        else:
            location_parts.append(f"pp. {pages}")
    location = ''.join(location_parts)
    
    doi_str = f"https://doi.org/{doi}" if doi else (url or "")
    
    # For APA: prevent double period (e.g. "Smith, J.. (2025)" → "Smith, J. (2025)")
    if style != "harvard" and author_str.endswith('.'):
        apa_author_str = author_str  # Already has period, don't add another
    elif style != "harvard":
        apa_author_str = author_str + '.'
    
    # ─── Harvard Style ───
    if style == "harvard":
        # Read additional metadata (gracefully skip if missing)
        editor = metadata.get("editor")
        edition = metadata.get("edition")
        place = metadata.get("place") or metadata.get("place_of_publication")
        accessed_date = metadata.get("accessed_date") or metadata.get("accessed")
        day_month = metadata.get("day_month")
        report_number = metadata.get("report_number")
        award = metadata.get("award")
        awarding_body = metadata.get("awarding_body")

        # Build publisher with optional place: "Place: Publisher"
        pub_str = f"{place}: {publisher}" if place and publisher else (publisher or place or "")

        # Build DOI/URL ending (no trailing full stop after DOI/URL)
        ending = ""
        if doi:
            ending = f" doi: https://doi.org/{doi}"
        elif url:
            ending = f" Available at: {url}"
            if accessed_date:
                ending += f" (Accessed: {accessed_date})"

        # ── Journal Article (Rules 4-6) ──
        if ref_type == "Journal Article" and source:
            ref_plain = f"{author_str} ({year}) '{title}', {source}"
            ref_html = f"{author_str_html} ({year}) '{title}', <i>{source}</i>"
            if location:
                ref_plain += f", {location}"
                ref_html += f", {location}"
            ref_plain += "."
            ref_html += "."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Book Chapter (Rule 3) ──
        elif ref_type == "Book Chapter":
            ref_plain = f"{author_str} ({year}) '{title}'"
            ref_html = f"{author_str_html} ({year}) '{title}'"
            if editor:
                ref_plain += f", in {editor}"
                ref_html += f", in {editor}"
            if source:  # Book title
                ref_plain += f" {source}"
                ref_html += f" <i>{source}</i>"
            if edition:
                ref_plain += f". {edition}"
                ref_html += f". {edition}"
            if pub_str:
                ref_plain += f". {pub_str}"
                ref_html += f". {pub_str}"
            if pages:
                ref_plain += f", pp. {pages}"
                ref_html += f", pp. {pages}"
            ref_plain += "."
            ref_html += "."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Book (Rules 1-2) ──
        elif ref_type == "Book":
            ref_plain = f"{author_str} ({year}) {title}."
            ref_html = f"{author_str_html} ({year}) <i>{title}</i>."
            if edition:
                ref_plain += f" {edition}."
                ref_html += f" {edition}."
            if pub_str:
                ref_plain += f" {pub_str}."
                ref_html += f" {pub_str}."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Report (Rules 8-9) ──
        elif ref_type == "Report":
            ref_plain = f"{author_str} ({year}) {title}."
            ref_html = f"{author_str_html} ({year}) <i>{title}</i>."
            if report_number:
                ref_plain += f" {report_number}."
                ref_html += f" {report_number}."
            if pub_str:
                ref_plain += f" {pub_str}."
                ref_html += f" {pub_str}."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Conference Paper (Rule 12) ──
        elif ref_type == "Conference Paper":
            ref_plain = f"{author_str} ({year}) '{title}'"
            ref_html = f"{author_str_html} ({year}) '{title}'"
            if source:  # Conference title
                ref_plain += f", {source}"
                ref_html += f", <i>{source}</i>"
            ref_plain += "."
            ref_html += "."
            if pub_str:
                ref_plain += f" {pub_str}"
                ref_html += f" {pub_str}"
            if pages:
                ref_plain += f", pp. {pages}"
                ref_html += f", pp. {pages}"
            ref_plain += "."
            ref_html += "."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Dissertation / Thesis (Rule 13) ──
        elif ref_type in ("Dissertation", "Thesis"):
            ref_plain = f"{author_str} ({year}) {title}."
            ref_html = f"{author_str_html} ({year}) <i>{title}</i>."
            if award:
                ref_plain += f" {award}."
                ref_html += f" {award}."
            if awarding_body:
                ref_plain += f" {awarding_body}."
                ref_html += f" {awarding_body}."
            elif publisher:
                ref_plain += f" {publisher}."
                ref_html += f" {publisher}."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Newspaper Article (Rule 14) ──
        elif ref_type == "Newspaper Article":
            ref_plain = f"{author_str} ({year}) '{title}'"
            ref_html = f"{author_str_html} ({year}) '{title}'"
            if source:  # Newspaper name
                ref_plain += f", {source}"
                ref_html += f", <i>{source}</i>"
            if day_month:
                ref_plain += f", {day_month}"
                ref_html += f", {day_month}"
            if pages:
                ref_plain += f", p. {pages}"
                ref_html += f", p. {pages}"
            ref_plain += "."
            ref_html += "."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Blog Post (Rule 10) ──
        elif ref_type == "Blog Post":
            ref_plain = f"{author_str} ({year}) '{title}'"
            ref_html = f"{author_str_html} ({year}) '{title}'"
            if source:  # Blog title
                ref_plain += f", {source}"
                ref_html += f", <i>{source}</i>"
            if day_month:
                ref_plain += f", {day_month}"
                ref_html += f", {day_month}"
            ref_plain += "."
            ref_html += "."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Podcast (Rule 15) ──
        elif ref_type == "Podcast":
            ref_plain = f"{author_str} ({year}) {title} [Podcast]."
            ref_html = f"{author_str_html} ({year}) <i>{title}</i> [Podcast]."
            if day_month:
                ref_plain += f" {day_month}."
                ref_html += f" {day_month}."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Dataset (Rule 18) ──
        elif ref_type == "Dataset":
            ref_plain = f"{author_str} ({year}) '{title}'."
            ref_html = f"{author_str_html} ({year}) '{title}'."
            if edition:
                ref_plain += f" {edition}."
                ref_html += f" {edition}."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Web Page (Rule 7) ──
        elif ref_type == "Web Page":
            ref_plain = f"{author_str} ({year}) {title}."
            ref_html = f"{author_str_html} ({year}) <i>{title}</i>."
            if ending:
                ref_plain += ending
                ref_html += ending

        # ── Fallback (unknown type) ──
        else:
            ref_plain = f"{author_str} ({year}) {title}."
            ref_html = f"{author_str_html} ({year}) <i>{title}</i>."
            if source:
                ref_plain += f" {source}."
                ref_html += f" <i>{source}</i>."
            if pub_str:
                ref_plain += f" {pub_str}."
                ref_html += f" {pub_str}."
            if ending:
                ref_plain += ending
                ref_html += ending
    
    # ─── APA 7th Style ───
    elif style == "apa":
        if ref_type == "Journal Article" and source:
            # Author, A. A. (Year). Title of article. Title of Periodical, vol(issue), pages. DOI
            ref_plain = f"{apa_author_str} ({year}). {title}. {source}"
            ref_html = f"{apa_author_str} ({year}). {title}. <i>{source}</i>"
            if location:
                if volume:
                    # Italicize volume number with source
                    ref_plain += f", {location}"
                    # In APA, the volume is also italicized (as part of journal name style)
                    if volume and issue:
                        ref_html += f", <i>{volume}</i>({issue})"
                        if pages:
                            ref_html += f", {pages}"
                            ref_plain_end = ""
                    elif volume:
                        ref_html += f", <i>{volume}</i>"
                        if pages:
                            ref_html += f", {pages}"
                    else:
                        ref_html += f", {location}"
                else:
                    ref_plain += f", {location}"
                    ref_html += f", {location}"
            ref_plain += "."
            ref_html += "."
            if doi_str:
                ref_plain += f" {doi_str}"
                ref_html += f" {doi_str}"
        elif ref_type in ("Book", "Book Chapter"):
            # Author, A. A. (Year). Title of work. Publisher. DOI
            ref_plain = f"{apa_author_str} ({year}). {title}."
            ref_html = f"{apa_author_str} ({year}). <i>{title}</i>."
            if publisher:
                ref_plain += f" {publisher}."
                ref_html += f" {publisher}."
            if doi_str:
                ref_plain += f" {doi_str}"
                ref_html += f" {doi_str}"
        elif ref_type == "Web Page":
            # Author, A. A. (Year). Title of work. Site Name. URL
            ref_plain = f"{apa_author_str} ({year}). {title}."
            ref_html = f"{apa_author_str} ({year}). <i>{title}</i>."
            if source:
                ref_plain += f" {source}."
                ref_html += f" {source}."
            if doi_str:
                ref_plain += f" {doi_str}"
                ref_html += f" {doi_str}"
        else:
            ref_plain = f"{apa_author_str} ({year}). {title}."
            ref_html = f"{apa_author_str} ({year}). <i>{title}</i>."
            if source:
                ref_plain += f" {source}."
                ref_html += f" <i>{source}</i>."
            if doi_str:
                ref_plain += f" {doi_str}"
                ref_html += f" {doi_str}"
    
    # ─── Vancouver (NLM) Style ───
    elif style == "vancouver":
        language = metadata.get("language")
        part_name = metadata.get("part_name")
        part_title = metadata.get("part_title")
        day_month = metadata.get("day_month")
        
        # Build Notes section (DOI / URL)
        notes = ""
        if doi:
            notes = f" doi:{doi}."
        elif url:
            notes = f" Available from: {url}"
            
        if ref_type == "Journal Article" and source:
            # Date builder: YYYY Mmm DD
            date_str = f"{year} {day_month}" if day_month else year
            
            # Location builder: Vol(Issue) / condensed pages
            van_loc_parts = []
            if volume and issue:
                van_loc_parts.append(f"{volume}({issue})")
            elif volume:
                van_loc_parts.append(volume)
                
            condensed_pages = condense_pages(pages) if pages else ""
            
            # Punctuation Chain Builder
            chain = f" {date_str}"
            if van_loc_parts:
                chain += f";{''.join(van_loc_parts)}"
                if condensed_pages:
                    chain += f":{condensed_pages}."
                else:
                    chain += "."
            else:
                # No volume/issue; separate date from pages with a colon
                if condensed_pages:
                    chain += f":{condensed_pages}."
                else:
                    chain += "."
            
            # ── Non-English Journal Article (Rule 3) ──
            if language and language.lower() != "english":
                ref_plain = f"{author_str} [{title}]. {source}."
                ref_html = f"{author_str} [{title}]. {source}."
                ref_plain += chain + f" {language}."
                ref_html += chain + f" {language}."
            
            # ── Part of a Journal Article (Rule 5) ──
            elif part_name and part_title and pages:
                # Build standard first without pages, then add part details
                base_chain = f" {date_str}"
                if van_loc_parts:
                    base_chain += f";{''.join(van_loc_parts)}"
                else:
                    base_chain += ":"
                base_chain += "."
                
                ref_plain = f"{author_str} {title}. {source}.{base_chain} {part_name}, {part_title}; p. {pages}."
                ref_html = f"{author_str} {title}. {source}.{base_chain} {part_name}, {part_title}; p. {pages}."
                
            # ── Standard Journal Article (Rules 1, 2, 4) ──
            else:
                ref_plain = f"{author_str} {title}. {source}.{chain}"
                ref_html = f"{author_str} {title}. {source}.{chain}"

            if notes:
                ref_plain += notes
                ref_html += notes

        # ── Book Chapter ──
        elif ref_type == "Book Chapter":
            # Chapter Author. Chapter title. In: Editor(s), editors. Book title. Place: Publisher; Year. p. Pages.
            editor = metadata.get("editor")
            place = metadata.get("place") or metadata.get("place_of_publication")
            
            ref_plain = f"{author_str} {title}."
            ref_html = f"{author_str} {title}."
            
            if editor:
                ref_plain += f" In: {editor}, editors."
                ref_html += f" In: {editor}, editors."
                
            if source: # Book Title
                ref_plain += f" {source}."
                ref_html += f" {source}."
                
            if place and publisher:
                ref_plain += f" {place}: {publisher};"
                ref_html += f" {place}: {publisher};"
            elif publisher:
                ref_plain += f" {publisher};"
                ref_html += f" {publisher};"
                
            ref_plain += f" {year}."
            ref_html += f" {year}."
            
            if pages:
                condensed = condense_pages(pages)
                ref_plain += f" p. {condensed}."
                ref_html += f" p. {condensed}."
                
            if notes:
                ref_plain += notes
                ref_html += notes

        # ── Book ──
        elif ref_type == "Book":
            # Author(s). Title. Edition. Place: Publisher; Year. Pages p.
            place = metadata.get("place") or metadata.get("place_of_publication")
            
            ref_plain = f"{author_str} {title}."
            ref_html = f"{author_str} {title}."
            
            if place and publisher:
                ref_plain += f" {place}: {publisher};"
                ref_html += f" {place}: {publisher};"
            elif publisher:
                ref_plain += f" {publisher};"
                ref_html += f" {publisher};"
                
            ref_plain += f" {year}."
            ref_html += f" {year}."
            
            if notes:
                ref_plain += notes
                ref_html += notes

        # ── Web Page ──
        elif ref_type == "Web Page":
            # Author. Title [Internet]. Place: Publisher; Year. Available from: URL
            place = metadata.get("place") or metadata.get("place_of_publication")
            
            ref_plain = f"{author_str} {title} [Internet]."
            ref_html = f"{author_str} {title} [Internet]."
            
            if place and publisher:
                ref_plain += f" {place}: {publisher};"
                ref_html += f" {place}: {publisher};"
            elif publisher:
                ref_plain += f" {publisher};"
                ref_html += f" {publisher};"
                
            ref_plain += f" {year}."
            ref_html += f" {year}."
            
            if notes:
                ref_plain += notes
                ref_html += notes

        # ── Dissertation / Thesis ──
        elif ref_type in ("Dissertation", "Thesis"):
            # Author. Title [dissertation]. Place: Publisher; Year.
            place = metadata.get("place") or metadata.get("place_of_publication")
            
            ref_plain = f"{author_str} {title} [dissertation]."
            ref_html = f"{author_str} {title} [dissertation]."
            
            if place and publisher:
                ref_plain += f" {place}: {publisher};"
                ref_html += f" {place}: {publisher};"
            elif publisher:
                ref_plain += f" {publisher};"
                ref_html += f" {publisher};"
                
            ref_plain += f" {year}."
            ref_html += f" {year}."
            
            if notes:
                ref_plain += notes
                ref_html += notes

        # ── Fallback ──
        else:
            ref_plain = f"{author_str} {title}."
            ref_html = f"{author_str} {title}."
            if source:
                ref_plain += f" {source}."
                ref_html += f" {source}."
            ref_plain += f" {year}."
            ref_html += f" {year}."
            
            if notes:
                ref_plain += notes
                ref_html += notes
    
    # ─── Fallback (unknown style, default to plain) ───
    else:
        ref_plain = f"{author_str} ({year}) {title}."
        ref_html = f"{author_str} ({year}) <i>{title}</i>."
        if source:
            ref_plain += f" {source}."
            ref_html += f" <i>{source}</i>."
        if doi_str:
            ref_plain += f" {doi_str}"
            ref_html += f" {doi_str}"
    
    return {
        "formatted": ref_plain,
        "formatted_html": ref_html,
        "type": ref_type,
        "metadata": metadata,
    }


@app.post("/api/extract-reference")
async def extract_reference(file: UploadFile = File(...), style: str = "harvard"):
    """Upload a PDF/DOCX and get a formatted reference for citing it."""
    if not file.filename.lower().endswith(('.pdf', '.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, DOCX, or DOC.")
    
    file_bytes = await file.read()
    magic = file_bytes[:8]
    
    try:
        if magic[:4] == b'%PDF':
            metadata = await extract_pdf_metadata(file_bytes)
        elif magic[:2] == b'PK':
            # DOCX: extract metadata from document properties
            metadata = extract_docx_metadata(file_bytes)
        elif magic[:4] == b'\xd0\xcf\x11\xe0':
            # DOC: limited metadata, extract text
            text = extract_doc_text(file_bytes)
            metadata = {
                "authors": None, "title": file.filename.rsplit('.', 1)[0],
                "year": None, "source": None, "doi": None, "url": None,
                "volume": None, "issue": None, "pages": None,
                "publisher": None, "type": "Other",
            }
            # Try to find DOI in text
            if text:
                doi_match = re.search(r'(?:doi[:\s]*|https?://(?:dx\.)?doi\.org/)(\S+?)(?:\s|$)', text, re.IGNORECASE)
                if doi_match:
                    metadata["doi"] = doi_match.group(1).rstrip('.,;)')
                year_match = re.search(r'\b(19|20)\d{2}\b', text[:2000])
                if year_match:
                    metadata["year"] = year_match.group(0)
        else:
            raise HTTPException(status_code=400, detail="Unrecognized file format.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read file: {str(e)}")
    
    result = format_reference(metadata, style)
    return result


@app.post("/api/reformat-reference")
async def reformat_reference(request: Request, style: str = "harvard"):
    """Reformat an already-extracted reference with a different style. No re-extraction needed."""
    body = await request.json()
    metadata = body.get("metadata")
    if not metadata:
        raise HTTPException(status_code=400, detail="Missing metadata in request body.")
    result = format_reference(metadata, style)
    return result


def extract_docx_metadata(file_bytes: bytes) -> dict:
    """Extract metadata from a DOCX file using document core properties."""
    metadata = {
        "authors": None, "title": None, "year": None,
        "source": None, "doi": None, "url": None,
        "volume": None, "issue": None, "pages": None,
        "publisher": None, "type": "Other",
    }
    
    with io.BytesIO(file_bytes) as f:
        doc = Document(f)
        
        # Core properties
        props = doc.core_properties
        if props.author:
            metadata["authors"] = props.author
        if props.title and len(props.title) > 3:
            metadata["title"] = props.title
        if props.created:
            metadata["year"] = str(props.created.year)
        
        # Parse text for DOI / URL
        full_text = '\n'.join(p.text for p in doc.paragraphs[:5])
        doi_match = re.search(r'(?:doi[:\s]*|https?://(?:dx\.)?doi\.org/)(\S+?)(?:\s|$)', full_text, re.IGNORECASE)
        if doi_match:
            metadata["doi"] = doi_match.group(1).rstrip('.,;)')
        
        if not metadata["title"]:
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if lines:
                metadata["title"] = lines[0]
    
    # CrossRef DOI lookup
    if metadata["doi"]:
        try:
            crossref_url = f"https://api.crossref.org/works/{metadata['doi']}"
            resp = requests.get(crossref_url, timeout=10, headers={
                'User-Agent': 'WritingTools/1.0 (mailto:support@paradoxlabs.com)'
            })
            if resp.status_code == 200:
                data = resp.json().get('message', {})
                
                authors = data.get('author', [])
                if authors:
                    author_parts = []
                    for a in authors:
                        family = a.get('family', '')
                        given = a.get('given', '')
                        if family and given:
                            initials = '. '.join(w[0].upper() for w in given.split() if w) + '.'
                            author_parts.append(f"{family}, {initials}")
                        elif family:
                            author_parts.append(family)
                    if author_parts:
                        metadata["authors"] = author_parts
                
                titles = data.get('title', [])
                if titles:
                    metadata["title"] = titles[0]
                
                date_parts = data.get('published-print', data.get('published-online', data.get('created', {})))
                if date_parts and 'date-parts' in date_parts:
                    parts = date_parts['date-parts'][0]
                    if parts:
                        metadata["year"] = str(parts[0])
                
                container = data.get('container-title', [])
                if container:
                    metadata["source"] = container[0]
                
                if data.get('volume'): metadata["volume"] = data['volume']
                if data.get('issue'): metadata["issue"] = data['issue']
                if data.get('page'): metadata["pages"] = data['page']
                if data.get('publisher'): metadata["publisher"] = data['publisher']
                
                cr_type = data.get('type', '')
                type_map = {
                    'journal-article': 'Journal Article',
                    'book': 'Book',
                    'book-chapter': 'Book Chapter',
                    'proceedings-article': 'Conference Paper',
                    'report': 'Report',
                }
                metadata["type"] = type_map.get(cr_type, 'Other')
        except Exception:
            pass
    
    if metadata["type"] == "Other":
        metadata["type"] = classify_source_type(metadata)
    
    if isinstance(metadata["authors"], str):
        raw = metadata["authors"]
        if '; ' in raw:
            metadata["authors"] = [a.strip() for a in raw.split(';') if a.strip()]
        elif ' and ' in raw or ' & ' in raw:
            raw = raw.replace(' & ', ' and ')
            metadata["authors"] = [a.strip() for a in raw.split(' and ') if a.strip()]
        else:
            metadata["authors"] = [raw.strip()]
    
    return metadata


class FormatRequest(BaseModel):
    references: List[str]
    style: str = "harvard"


def _value_in_input(value: str, ref_text: str, threshold: float = 0.6) -> bool:
    """
    Anti-hallucination check: verify that a value extracted by AI actually
    appears in the original input text. Uses substring match first (fast),
    then falls back to fuzzy similarity for minor reformatting differences
    (e.g. "Smith, J." vs "Smith J").
    """
    if not value or not ref_text:
        return False
    val_lower = value.lower().strip()
    ref_lower = ref_text.lower()
    # Fast path: exact substring
    if val_lower in ref_lower:
        return True
    # Strip punctuation and check again (handles "Smith, J." vs "Smith J")
    val_clean = re.sub(r'[.,;:\'"()\[\]{}]', '', val_lower).strip()
    ref_clean = re.sub(r'[.,;:\'"()\[\]{}]', '', ref_lower)
    if val_clean in ref_clean:
        return True
    # Fuzzy match: the value's words should mostly appear in the input
    ratio = difflib.SequenceMatcher(None, val_clean, ref_clean).ratio()
    # For short values (like years), require near-exact match
    if len(val_clean) <= 6:
        return val_clean in ref_clean
    return ratio >= threshold


async def parse_raw_reference(ref_text: str) -> dict:
    """
    Parse a raw reference string into structured metadata.
    Pipeline: Regex first → AI extraction with Input Containment Check.
    
    Anti-hallucination strategy: The user says the pasted reference is already
    correct. Therefore every field AI extracts MUST be traceable back to the
    original input text. If AI returns a value not found in the input, it is
    rejected as hallucinated — no external API needed.
    """
    metadata = {
        "authors": None, "title": None, "year": None,
        "source": None, "doi": None, "url": None,
        "volume": None, "issue": None, "pages": None,
        "publisher": None, "type": "Other",
    }
    field_sources = {}

    # ─── Step 1: Regex extraction from raw text (zero hallucination risk) ───
    # DOI
    doi_match = re.search(
        r'(?:doi[:\s]*|https?://(?:dx\.)?doi\.org/)(10\.\d{4,}/[a-zA-Z0-9.\-_/:()]+)',
        ref_text, re.IGNORECASE
    )
    if doi_match:
        doi = doi_match.group(1).rstrip('.,;)')
        metadata["doi"] = doi
        field_sources["doi"] = "text_parsing"

    # Year
    year_match = re.search(r'\b((?:19|20)\d{2})\b', ref_text)
    if year_match:
        metadata["year"] = year_match.group(1)
        field_sources["year"] = "text_parsing"

    # Volume/Issue
    vol_match = re.search(r'(\d+)\s*\((\d+)\)', ref_text)
    if vol_match:
        metadata["volume"] = vol_match.group(1)
        metadata["issue"] = vol_match.group(2)
        field_sources["volume"] = "text_parsing"
        field_sources["issue"] = "text_parsing"

    # Pages
    pages_match = re.search(r'(?:pp?\.?\s*)?(\d+)\s*[-–]\s*(\d+)', ref_text)
    if pages_match:
        metadata["pages"] = f"{pages_match.group(1)}-{pages_match.group(2)}"
        field_sources["pages"] = "text_parsing"

    # URL (if no DOI)
    if not metadata["doi"]:
        url_match = re.search(r'(https?://\S+)', ref_text)
        if url_match:
            metadata["url"] = url_match.group(1).rstrip('.,;)')
            field_sources["url"] = "text_parsing"

    # ─── Step 2: AI extraction with Input Containment Check ───
    # AI identifies which parts of the text are authors, title, source, etc.
    # Every value is then verified against the original input — if it's not
    # found in the input, it's rejected as hallucinated.
    try:
        model_name = 'gemini-3-flash-preview'
        client = get_client(model=model_name)

        prompt = f"""You are a metadata extraction tool. Extract ONLY what you can see in the reference text below.
DO NOT invent, guess, or hallucinate any information. If a field is not clearly visible, return null.
Every value you return MUST come directly from the text — do not rephrase, infer, or add information.

Extract these fields from the reference:
- authors: List of author names exactly as they appear in the text (e.g. ["Clair, A.", "Hughes, A."])
- title: The main title of the work, exactly as written in the text
- year: The publication year as it appears in the text
- source: Journal name or publisher, exactly as written in the text
- doi: The DOI if present, exactly as written

REFERENCE TEXT:
{ref_text}

Respond in strict JSON only:
{{
    "authors": ["author1", "author2"] or null,
    "title": "title text" or null,
    "year": "2025" or null,
    "source": "journal or publisher name" or null,
    "doi": "10.1234/example" or null
}}"""

        response = await gemini_request_with_retry(client, prompt, model=model_name)
        ai_text = response.text.strip()
        if ai_text.startswith('```json'):
            ai_text = ai_text[7:].strip()
        if ai_text.endswith('```'):
            ai_text = ai_text[:-3].strip()
        ai_data = json.loads(ai_text)
        print(f"[Formatter AI] Raw extraction: {ai_data}")

        # ─── Input Containment Check: reject anything not in the original text ───
        # Title
        ai_title = ai_data.get("title")
        if ai_title and _value_in_input(ai_title, ref_text):
            if not metadata.get("title"):
                metadata["title"] = ai_title
                field_sources["title"] = "ai_verified"
        elif ai_title:
            print(f"[Formatter Anti-Hallucination] REJECTED title: '{ai_title}' — not found in input")

        # Authors
        ai_authors = ai_data.get("authors")
        if ai_authors and isinstance(ai_authors, list):
            verified_authors = []
            for author in ai_authors:
                if _value_in_input(author, ref_text):
                    verified_authors.append(author)
                else:
                    # Check if at least the surname is in the input
                    surname = author.split(',')[0].strip() if ',' in author else author.split()[0].strip()
                    if _value_in_input(surname, ref_text):
                        verified_authors.append(author)
                    else:
                        print(f"[Formatter Anti-Hallucination] REJECTED author: '{author}' — not found in input")
            if verified_authors:
                metadata["authors"] = verified_authors
                field_sources["authors"] = "ai_verified"

        # Source (journal/publisher)
        ai_source = ai_data.get("source")
        if ai_source and _value_in_input(ai_source, ref_text):
            if not metadata.get("source"):
                metadata["source"] = ai_source
                field_sources["source"] = "ai_verified"
        elif ai_source:
            print(f"[Formatter Anti-Hallucination] REJECTED source: '{ai_source}' — not found in input")

        # Year — only accept if it matches what regex found or is in the text
        ai_year = ai_data.get("year")
        if ai_year and str(ai_year) in ref_text:
            if not metadata.get("year"):
                metadata["year"] = str(ai_year)
                field_sources["year"] = "ai_verified"

        # DOI — only accept if it's actually in the text
        ai_doi = ai_data.get("doi")
        if ai_doi and ai_doi in ref_text and not metadata.get("doi"):
            metadata["doi"] = ai_doi
            field_sources["doi"] = "ai_verified"

    except Exception as e:
        print(f"[Formatter AI] FAILED: {type(e).__name__}: {e}")

    # ─── Step 3: Classify type ───
    if metadata["type"] == "Other":
        metadata["type"] = classify_source_type(metadata)

    # Ensure authors is a list
    if isinstance(metadata["authors"], str):
        raw = metadata["authors"]
        if '; ' in raw:
            metadata["authors"] = [a.strip() for a in raw.split(';') if a.strip()]
        elif ' and ' in raw or ' & ' in raw:
            raw = raw.replace(' & ', ' and ')
            metadata["authors"] = [a.strip() for a in raw.split(' and ') if a.strip()]
        else:
            metadata["authors"] = [raw.strip()]

    metadata["field_sources"] = field_sources
    return metadata


@app.post("/api/format")
async def format_references(req: FormatRequest):
    """
    Parse raw reference text into metadata, then format deterministically.
    Returns metadata alongside formatted output so the frontend can do instant style switching.
    """
    formatted_refs = []
    for ref in req.references:
        try:
            metadata = await parse_raw_reference(ref)
            result = format_reference(metadata, req.style)
            # Include the original raw text
            result["original"] = ref
            formatted_refs.append(result)
        except Exception as e:
            formatted_refs.append({"original": ref, "error": str(e)})

    return {"formatted_references": formatted_refs}

@app.get("/api-key-usage")
async def api_key_usage():
    """Return current API key usage statistics for frontend dashboard."""
    if KEY_MANAGER_AVAILABLE:
        manager = get_api_key_manager()
        return manager.get_status()
    return {
        "service": "Google",
        "total_keys": 1 if API_KEY else 0,
        "available_keys": 1 if API_KEY else 0,
        "exhausted_keys": 0,
        "keys": [],
        "message": "Key manager not available, using single key mode"
    }

# ─── Full-Text Search Endpoints ───

@app.post("/api/search/upload")
async def search_upload(files: list[UploadFile] = File(...)):
    """Upload PDF(s) and index their full text for searching."""
    indexed = []
    errors = []
    for file in files:
        if not file.filename.lower().endswith('.pdf'):
            errors.append({"filename": file.filename, "error": "Only PDF files are supported"})
            continue
        try:
            file_bytes = await file.read()
            if file_bytes[:4] != b'%PDF':
                errors.append({"filename": file.filename, "error": "Not a valid PDF file"})
                continue
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                text = page.extract_text() or ''
                pages_text.append(text)
            if not any(t.strip() for t in pages_text):
                errors.append({"filename": file.filename, "error": "No text could be extracted from PDF"})
                continue
            doc_id = search_store.index_document(file.filename, pages_text)
            indexed.append({
                "doc_id": doc_id,
                "filename": file.filename,
                "total_pages": len(pages_text),
            })
        except Exception as e:
            errors.append({"filename": file.filename, "error": str(e)})
    return {"indexed": indexed, "errors": errors}


class SearchQuery(BaseModel):
    queries: list[str]  # Multiple queries to search simultaneously

@app.post("/api/search/query")
async def search_query(req: SearchQuery):
    """Search for multiple phrases across all indexed documents simultaneously."""
    # Filter out empty queries
    clean_queries = [q.strip() for q in req.queries if q.strip()]
    if not clean_queries:
        raise HTTPException(status_code=400, detail="At least one query is required")
    
    all_results = []
    for query_text in clean_queries:
        matches = search_store.search(query_text)
        all_results.append({
            "query": query_text,
            "results": matches,
            "total": len(matches),
        })
    return {"groups": all_results, "total_queries": len(clean_queries)}


@app.delete("/api/search/document/{doc_id}")
async def search_delete_document(doc_id: str):
    """Remove a document from the search index."""
    success = search_store.delete_document(doc_id)
    if not success:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"deleted": doc_id}


@app.get("/api/search/documents")
async def search_list_documents():
    """List all indexed documents."""
    docs = search_store.list_documents()
    return {"documents": docs}


# ─── Humanizer (Cognitive Synthesizer) Endpoints ───

import humanizer_store
import humanizer


class HumanizeRequest(BaseModel):
    text: str


@app.post("/api/humanizer/upload")
async def humanizer_upload(files: list[UploadFile] = File(...)):
    """Upload PDF(s) and mine their sentences into the skeleton bank using LLM."""
    indexed = []
    errors = []
    for file in files:
        if not file.filename.lower().endswith('.pdf'):
            errors.append({"filename": file.filename, "error": "Only PDF files are supported"})
            continue
        try:
            file_bytes = await file.read()
            if file_bytes[:4] != b'%PDF':
                errors.append({"filename": file.filename, "error": "Not a valid PDF file"})
                continue
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                text = page.extract_text() or ''
                pages_text.append(text)
            if not any(t.strip() for t in pages_text):
                errors.append({"filename": file.filename, "error": "No text could be extracted from PDF"})
                continue
            result = await humanizer_store.index_document(file.filename, pages_text)
            indexed.append({
                "doc_id": result["doc_id"],
                "filename": file.filename,
                "skeleton_count": result["skeleton_count"],
                # Keep backward compat
                "sentence_count": result["skeleton_count"],
            })
        except Exception as e:
            errors.append({"filename": file.filename, "error": str(e)})
    return {"indexed": indexed, "errors": errors}


@app.post("/api/humanizer/humanize")
async def humanizer_humanize(req: HumanizeRequest):
    """Run the Cognitive Synthesizer pipeline on the given text."""
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="Text is required")
    
    # Check if there are any skeletons in the bank
    stats = humanizer_store.get_stats()
    if stats["total_skeletons"] == 0:
        raise HTTPException(
            status_code=400,
            detail="No human skeletons in database. Upload PDFs with human-written text first."
        )
    
    result = await humanizer.humanize_text(req.text)
    return result


@app.get("/api/humanizer/documents")
async def humanizer_list_documents():
    """List all documents in the skeleton bank."""
    docs = humanizer_store.list_documents()
    return {"documents": docs}


@app.delete("/api/humanizer/document/{doc_id}")
async def humanizer_delete_document(doc_id: str):
    """Remove a document and its skeletons from the bank."""
    success = humanizer_store.delete_document(doc_id)
    if not success:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"deleted": doc_id}


@app.get("/api/humanizer/stats")
async def humanizer_stats():
    """Get stats about the skeleton bank."""
    stats = humanizer_store.get_stats()
    return stats


@app.get("/api/health")
async def health_check():
    return {"status": "ok"}


from pydantic import BaseModel
class PhrasebankRequest(BaseModel):
    text: str

@app.post("/api/phrasebank/rewrite")
async def api_phrasebank_rewrite(req: PhrasebankRequest):
    """Rewrite a sentence using Academic Phrasebank principles."""
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="Text is required")

    from phrasebank import process_phrasebank_rewrite
    result = await process_phrasebank_rewrite(req.text)
    if "error" in result:
        raise HTTPException(status_code=500, detail=result["error"])
         
    return result

import phrasebank_store

@app.post("/api/phrasebank/upload")
async def phrasebank_upload(files: list[UploadFile] = File(...)):
    """Upload PDF(s) and mine them for academic phrase templates."""
    indexed = []
    errors = []
    for file in files:
        if not file.filename.lower().endswith('.pdf'):
            errors.append({"filename": file.filename, "error": "Only PDF files are supported"})
            continue
        try:
            file_bytes = await file.read()
            if file_bytes[:4] != b'%PDF':
                errors.append({"filename": file.filename, "error": "Not a valid PDF file"})
                continue
            
            # Using the same PDF extractor from main.py
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                text = page.extract_text() or ''
                pages_text.append(text)
                
            if not any(t.strip() for t in pages_text):
                errors.append({"filename": file.filename, "error": "No text could be extracted from PDF"})
                continue
                
            result = await phrasebank_store.index_document(file.filename, pages_text)
            indexed.append({
                "doc_id": result["doc_id"],
                "filename": file.filename,
                "phrase_count": result["phrase_count"],
            })
        except Exception as e:
            errors.append({"filename": file.filename, "error": str(e)})
    return {"indexed": indexed, "errors": errors}

@app.get("/api/phrasebank/documents")
async def phrasebank_list_documents():
    """List all mapped phrasebank documents."""
    docs = phrasebank_store.list_documents()
    return {"documents": docs}

@app.delete("/api/phrasebank/document/{doc_id}")
async def phrasebank_delete_document(doc_id: str):
    """Remove a document and its phrases from the Phrasebank."""
    success = phrasebank_store.delete_document(doc_id)
    if not success:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"deleted": doc_id}

@app.get("/api/phrasebank/stats")
async def phrasebank_stats():
    """Get stats about the phrasebank database."""
    stats = phrasebank_store.get_stats()
    return stats

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
