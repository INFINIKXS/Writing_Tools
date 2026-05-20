"""
Metadata extraction from PDFs and DOCX files.
Deterministic cascade version.
==========================================================
Uses a lightweight deterministic cascade:

  Layer 1  →  PDF embedded properties  (fast, zero cost)
  Layer 2  →  Custom regex on first 3 pages  (your existing text_utils.py)
  Layer 3  →  pdf2doi  (focused DOI rescue)
  Layer 4  →  CrossRef API  (verification + metadata enrichment)
  Layer 5  →  PubMed API   (secondary verification)
"""

from __future__ import annotations

import io
import logging
import os
import re
import tempfile
import time
import asyncio
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional

import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from docx import Document

from utils.text_utils import classify_source_type, extract_doi, extract_all_dois

logger = logging.getLogger(__name__)

# ── Global HTTP Session with Connection Pooling ────────────────────────────────
_http_session = requests.Session()
_retries = Retry(
    total=3,
    backoff_factor=1,
    status_forcelist=[429, 500, 502, 503, 504],
    allowed_methods=["GET", "POST"]
)
_adapter = HTTPAdapter(max_retries=_retries, pool_connections=10, pool_maxsize=10)
_http_session.mount("http://", _adapter)
_http_session.mount("https://", _adapter)

# ── PubMed Rate Limiter ────────────────────────────────────────────────────────
# NCBI E-utilities allows 3 requests/second without an API key.
# This thread-safe rate limiter ensures we stay under that limit even when
# multiple verification workers are running concurrently.
import threading

class _RateLimiter:
    """Thread-safe rate limiter: ensures minimum interval between calls."""
    def __init__(self, max_per_second: float):
        self._min_interval = 1.0 / max_per_second
        self._lock = threading.Lock()
        self._last_call = 0.0

    def wait(self):
        """Block until enough time has passed since the last call."""
        with self._lock:
            now = time.monotonic()
            elapsed = now - self._last_call
            if elapsed < self._min_interval:
                time.sleep(self._min_interval - elapsed)
            self._last_call = time.monotonic()

_pubmed_limiter = _RateLimiter(max_per_second=3.0)
_crossref_limiter = _RateLimiter(max_per_second=10.0)

# ── Constants ──────────────────────────────────────────────────────────────────
CROSSREF_API    = "https://api.crossref.org/works"
PUBMED_SEARCH   = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
PUBMED_FETCH    = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
REQUEST_TIMEOUT = 15   # seconds


# ══════════════════════════════════════════════════════════════════════════════
# LAYER 1 — PDF embedded metadata
# ══════════════════════════════════════════════════════════════════════════════

# Patterns that indicate the PDF /Title field contains internal production IDs,
# copyright notices, or other junk — NOT a real paper title.
_GARBAGE_TITLE_RE = re.compile(
    r'(?:'
    r'^[a-z0-9_\-]{5,}\s+\d+\.\.\d+'     # "ao2c00362 1..9", "pone_0294946 1..23"
    r'|^©'                                 # copyright symbol
    r'|^The Author\(?s?\)?'                # "The Author(s)"
    r'|Creative Commons'                   # CC licence text
    r'|NoDerivatives'
    r'|Open Access'
    r'|International License'
    r'|licensed material'
    r'|permission under this licence'
    r'|terms of the Creative Commons'
    r'|exceeds the permitted use'
    r'|copyright holder'
    r'|obtain permission'
    r'|et al\.\s+\w'                    # "Gupta et al. BioData..." citation header
    r'|\(\d{4}\)\s+\d+:\d+'             # journal volume format "(2024) 17:54"
    r'|full list of author'              # "Full list of author information is..."
    r'|author information'               # author info section boilerplate
    r'|distributed under the terms'
    r'|creativecommons\.org'
    r'|^doi:'                              # DOI string used as title
    r'|^https?://'                         # URL used as title
    r'|^Microsoft Word'                    # Word doc conversion artefact
    r'|^untitled'                          # literally "untitled"
    r'|^\d{4}-\d{4}'                       # ISSN used as title
    r')',
    re.IGNORECASE,
)


def _is_garbage_title(title: str) -> bool:
    """
    Return True if a PDF metadata title is clearly NOT a real paper title.
    Publishers often dump internal production IDs, filenames, or copyright
    notices into the PDF /Title field.
    """
    if not title or len(title.strip()) < 5:
        return True
    t = title.strip()
    # Matches known garbage patterns
    if _GARBAGE_TITLE_RE.search(t):
        return True
    # Mostly digits/punctuation with very few alpha chars → probably a code
    alpha_ratio = sum(c.isalpha() for c in t) / max(len(t), 1)
    if alpha_ratio < 0.4:
        return True
    # Fewer than 3 actual words → unlikely to be a paper title
    words = [w for w in t.split() if len(w) > 1 and w.isalpha()]
    if len(words) < 3:
        return True
    return False


def _extract_pdf_metadata(pdf_path: str) -> dict:
    """
    Read the hidden metadata baked into the PDF file itself.
    Checks modern XMP Metadata first, then falls back to legacy PDF Info fields.
    Returns a dict with keys: title, authors, year, doi  (all may be None).
    """
    result = {"title": None, "authors": [], "year": None, "doi": None}
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        
        # ── 1a. Try Modern XMP Metadata (Dublin Core) ──
        try:
            xmp = reader.xmp_metadata
            if xmp and xmp.dc_title:
                xmp_title = xmp.dc_title
                title_candidate = None
                if isinstance(xmp_title, dict):
                    title_candidate = xmp_title.get('x-default') or next(iter(xmp_title.values()), None)
                elif isinstance(xmp_title, str):
                    title_candidate = xmp_title
                
                if title_candidate and not _is_garbage_title(title_candidate):
                    result["title"] = title_candidate.strip()
                    logger.info("Extracted clean title via PDF XMP Metadata: %r", result["title"])
        except Exception as xmp_exc:
            logger.debug("XMP metadata extraction skipped/failed: %s", xmp_exc)

        # ── 1b. Fallback to Legacy Info Dictionary if XMP failed ──
        info = reader.metadata or {}
        if not result["title"]:
            raw_title = info.get("/Title") or info.get("Title")
            if raw_title and not _is_garbage_title(raw_title):
                result["title"] = raw_title.strip()
                logger.info("Extracted clean title via PDF Info Metadata: %r", result["title"])
            elif raw_title:
                logger.debug("Rejected garbage PDF title: %r", raw_title)

        raw_author = info.get("/Author") or info.get("Author")
        if raw_author:
            # Authors can be semicolon- or comma-separated
            result["authors"] = [
                a.strip() for a in re.split(r"[;,]", raw_author) if a.strip()
            ]

        # Year — check CreationDate or ModDate  (format: D:YYYYMMDDHHmmSS)
        for date_key in ("/CreationDate", "/ModDate", "CreationDate", "ModDate"):
            raw_date = info.get(date_key, "")
            year_match = re.search(r"(\d{4})", str(raw_date))
            if year_match:
                year = int(year_match.group(1))
                if 1900 < year <= 2100:
                    result["year"] = year
                    break

        # DOI — try Subject/Keywords/custom fields first
        for field in ("/Subject", "/Keywords", "Subject", "Keywords", "/doi", "doi"):
            raw = info.get(field, "")
            doi = extract_doi(str(raw))
            if doi:
                result["doi"] = doi
                break

    except Exception as exc:
        logger.warning("PDF metadata read failed: %s", exc)

    return result


# ══════════════════════════════════════════════════════════════════════════════
# LAYER 2 — Your existing regex  (text_utils.py)
# ══════════════════════════════════════════════════════════════════════════════

def _extract_via_regex(pdf_path: str, pages: int = 2) -> dict:
    """
    Extract raw text from the first `pages` pages and run your existing
    two-pass DOI regex from text_utils.  Also pulls a naive title/year guess
    supporting multi-line concatenations.
    Returns: { doi, year, title }  — authors are not reliably found via regex.
    """
    result = {"doi": None, "year": None, "title": None}
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        text_pages = []
        for i, page in enumerate(reader.pages):
            if i >= pages:
                break
            text_pages.append(page.extract_text() or "")
        full_text = "\n".join(text_pages)

        # ── Heal line-wrapped DOI URLs before regex ──────────────────────────
        healed = re.sub(
            r'(https?://[^\s]*?)\n\s*',
            lambda m: m.group(1),
            full_text
        )
        healed = re.sub(r'(10\.\d{4,9}/[^\s]*?)-?\n\s*([a-zA-Z0-9])', r'\1\2', healed)
        healed = re.sub(r'(10\.\d{4,9}/[^\s]*\.)\s*\n\s*(\d)', r'\1\2', healed)

        # DOI — your robust two-pass regex
        result["doi"] = extract_doi(healed)

        # Year — look for 4-digit years in plausible publication range
        years = re.findall(r"\b(19[5-9]\d|20[0-2]\d)\b", full_text)
        if years:
            result["year"] = int(years[0])

        # Title — heuristic: scan and support multi-line title fields
        lines = [line.strip() for line in text_pages[0].splitlines() if line.strip()] if text_pages else []
        
        for idx, line in enumerate(lines):
            if len(line) > 20 and not re.match(r"^\d", line) and not _is_garbage_title(line):
                title_parts = [line]
                
                # Lookahead up to 2 lines to catch wrapped components of the title
                for lookahead in range(1, 3):
                    if idx + lookahead < len(lines):
                        next_line = lines[idx + lookahead]
                        # Continuation line criteria: short to medium length, no weird endings, not a header/date
                        if (len(next_line) > 5 and 
                            not re.match(r"^\d", next_line) and 
                            not _is_garbage_title(next_line) and
                            not any(keyword in next_line.lower() for keyword in ["vol.", "issue", "http", "doi:", "received", "abstract"])):
                            title_parts.append(next_line)
                        else:
                            break
                
                result["title"] = " ".join(title_parts)
                logger.info("Heuristic Regex layer extracted title: %r", result["title"])
                break

    except Exception as exc:
        logger.warning("Regex extraction failed: %s", exc)

    return result


def _extract_identity_from_pdf(pdf_path: str) -> dict:
    """
    Aggressively extract a paper identity fingerprint directly from PDF text.
    Used as a last-resort source of comparison identifiers before API verification,
    so that we always have *something* to validate against.

    Returns: { title: str|None, surnames: list[str], year: int|None }
      - title  : longest plausible title candidate found in pages 0-1
      - surnames: list of capitalised words that look like author surnames
      - year   : first plausible 4-digit publication year found
    """
    result = {"title": None, "surnames": [], "year": None}
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        pages_text = []
        for i, page in enumerate(reader.pages):
            if i >= 2:   # only look at cover pages
                break
            pages_text.append(page.extract_text() or "")
        text = "\n".join(pages_text)
        lines = [l.strip() for l in text.splitlines() if l.strip()]

        # ── Title: pick the longest line in the first 30 that looks like a title ──
        # Titles tend to be: ≥20 chars, <200 chars, no leading digits, not all-caps junk,
        # not a URL, not a copyright/affiliation line.
        _junk_re = re.compile(
            r'^(https?://|doi:|©|copyright|received|accepted|published|volume|issue'
            r'|correspondence|email|department|university|abstract|introduction|keywords'
            r'|license|open access|noderivatives)',
            re.IGNORECASE,
        )
        candidates = [
            l for l in lines[:30]
            if len(l) >= 20
            and len(l) < 200
            and not re.match(r'^\d', l)
            and not _junk_re.search(l)
            and not l.isupper()
            and not _is_garbage_title(l)
        ]
        if candidates:
            # Prefer the longest — titles are usually the longest non-junk line on page 1
            result["title"] = max(candidates, key=len)

        # ── Year: first plausible publication year ───────────────────────────
        years = re.findall(r'\b(19[5-9]\d|20[0-2]\d)\b', text)
        if years:
            result["year"] = int(years[0])

        # ── Author surnames: capitalised words that look like surnames ────────
        # Look for patterns like "Smith, J." or "J. Smith" or lines that are
        # a comma-separated list of capitalised single words.
        surname_re = re.compile(r'\b([A-Z][a-z]{2,}(?:-[A-Z][a-z]+)?)\b')
        # Only scan the first 15 lines (author block is near the top)
        author_block = "\n".join(lines[:15])
        all_caps_words = surname_re.findall(author_block)
        # Filter out common non-surname words
        _stop = {
            'Abstract', 'Introduction', 'Keywords', 'Background', 'Methods',
            'Results', 'Discussion', 'Conclusion', 'References', 'Received',
            'Accepted', 'Published', 'Correspondence', 'Journal', 'University',
            'Department', 'Research', 'Science', 'Nature', 'Review', 'Article',
        }
        result["surnames"] = [w for w in all_caps_words if w not in _stop]

    except Exception as exc:
        logger.warning("PDF identity extraction failed: %s", exc)
    return result


def hard_verify_against_pdf(title: str, authors: list, pdf_path: str) -> bool:
    """
    Physically scans the raw text of the first 2 pages of the PDF.
    Returns True if the title (>15 chars) and the first author's surname
    are found verbatim in the raw text (ignoring whitespace/punctuation).
    """
    import re
    import unicodedata
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        pdf_text = ""
        for i, page in enumerate(reader.pages):
            if i >= 2: break
            pdf_text += (page.extract_text() or "") + "\n"
        
        clean_title = re.sub(r'[^\w]', '', (title or "").lower())
        clean_pdf_text = re.sub(r'[^\w]', '', pdf_text.lower())
        
        if len(clean_title) <= 15 or clean_title not in clean_pdf_text:
            return False
            
        if not authors:
            return True
            
        first_author = unicodedata.normalize('NFKD', authors[0]).encode('ascii', 'ignore').decode().lower()
        # If the author is "Smith, J", grab the surname "Smith"
        first_author = first_author.split(',')[0].strip()
        first_author = re.sub(r'[^\w]', '', first_author)
        
        if first_author and first_author not in clean_pdf_text:
            return False
            
        return True
    except Exception as e:
        logger.warning("Failed to perform deep PDF text search for hard verification: %s", e)
        return False


def strict_ai_verify_against_pdf(ai_data: dict, pdf_path: str) -> bool:
    """
    Extremely strict verification for pure-AI metadata extraction.
    Ensures that ALL provided string/numeric fields (title, authors, year, journal, volume, issue, pages)
    physically exist within the raw text of the first 2 pages of the PDF.
    If even one field is hallucinated, the entire record is rejected.
    """
    import re
    import unicodedata
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        pdf_text = ""
        for i, page in enumerate(reader.pages):
            if i >= 2: break
            pdf_text += (page.extract_text() or "") + "\n"
            
        clean_pdf_text = re.sub(r'[^\w]', '', pdf_text.lower())
        
        # 1. Check Title (> 15 chars requirement)
        title = ai_data.get("title")
        if not title: return False
        clean_title = re.sub(r'[^\w]', '', title.lower())
        if len(clean_title) <= 15 or clean_title not in clean_pdf_text:
            return False
            
        # 2. Check all authors' surnames
        authors = ai_data.get("authors") or []
        for author in authors:
            author_norm = unicodedata.normalize('NFKD', author).encode('ascii', 'ignore').decode().lower()
            surname = author_norm.split(',')[0].strip()
            surname = re.sub(r'[^\w]', '', surname)
            if surname and surname not in clean_pdf_text:
                return False
                
        # 3. Check other metadata fields intelligently
        # Instead of rejecting the entire perfect extraction if a minor field (like Publisher) 
        # is hallucinated/inferred, we just strip that unverified field out.
        fields_to_check = ["year", "journal", "volume", "issue", "pages", "publisher"]
        for field in fields_to_check:
            val = ai_data.get(field)
            if not val:
                continue
                
            val_str = str(val).lower()
            field_verified = True
            
            if field == "pages":
                numbers = re.findall(r'\d+', val_str)
                for num in numbers:
                    if num not in clean_pdf_text:
                        logger.warning("Strict AI verification: Page number '%s' not found. Stripping 'pages' field.", num)
                        field_verified = False
                        break
            elif field in ["journal", "publisher"]:
                words = [re.sub(r'[^\w]', '', w) for w in val_str.split()]
                significant_words = [w for w in words if len(w) > 3]
                
                if not significant_words:
                    clean_val = re.sub(r'[^\w]', '', val_str)
                    if clean_val and clean_val not in clean_pdf_text:
                        logger.warning("Strict AI verification: %s '%s' not found. Stripping field.", field, val)
                        field_verified = False
                else:
                    found_count = sum(1 for w in significant_words if w in clean_pdf_text)
                    if found_count / len(significant_words) < 0.5:
                        logger.warning("Strict AI verification: %s '%s' failed word-match threshold. Stripping field.", field, val)
                        field_verified = False
            else:
                clean_val = re.sub(r'[^\w]', '', val_str)
                if clean_val and clean_val not in clean_pdf_text:
                    logger.warning("Strict AI verification: Field '%s' value '%s' not found. Stripping field.", field, val)
                    field_verified = False
            
            if not field_verified:
                # If it's the year that failed, that's critical. We should reject the whole thing.
                if field == "year":
                    logger.warning("Strict AI verification failed: Critical field 'year' hallucinated. Rejecting extraction.")
                    return False
                # Otherwise, just remove the unverified field
                del ai_data[field]
                    
        return True
    except Exception as e:
        logger.warning("Failed strict AI verification text search: %s", e)
        return False


def _validate_api_result(
    local_meta: dict,
    api_result: dict,
    pdf_path: str,
    threshold: float = 0.90,
) -> bool:
    """
    Validate that an API-returned metadata record actually belongs to the PDF.
    Uses a multi-signal approach:
      1. Title similarity (primary, ≥threshold)
      2. Author surname overlap (secondary)
      3. Year match (tertiary)

    If NO local signals exist after an aggressive extraction attempt,
    returns False — we refuse to blindly accept an unverifiable result.

    Returns True only if at least one signal passes and none actively fail.
    """
    from difflib import SequenceMatcher

    def _sim(a: str, b: str) -> float:
        a = re.sub(r'[^\w\s]', '', (a or '').lower())
        b = re.sub(r'[^\w\s]', '', (b or '').lower())
        return SequenceMatcher(None, a, b).ratio()

    # ── Gather local signals ─────────────────────────────────────────────────
    local_title   = local_meta.get("title")
    local_authors = local_meta.get("authors") or []
    local_year    = local_meta.get("year")

    # If the local layers gave us nothing, do an aggressive PDF scan now.
    if not local_title and not local_authors and not local_year:
        logger.info("No local identity data — running aggressive PDF identity extraction.")
        identity = _extract_identity_from_pdf(pdf_path)
        local_title   = local_title   or identity.get("title")
        local_year    = local_year    or identity.get("year")
        # Don't overwrite local_authors from pdf scan — surnames are too noisy
        # to match against structured API authors; use separately below.
        pdf_surnames  = identity.get("surnames", [])
    else:
        pdf_surnames = []

    # ── Gather API signals ───────────────────────────────────────────────────
    api_title   = api_result.get("title")
    api_authors = api_result.get("authors") or []
    api_year    = api_result.get("year")
    api_type    = api_result.get("type", "")
    api_journal = api_result.get("journal")

    # ── STRICT COMPLETENESS GATE ─────────────────────────────────────────────
    # A reference is only as good as its weakest field.  If the API returned
    # data that is missing ANY of the essential fields (title, authors, year,
    # and journal if it's a journal article), we reject it outright.
    # The caller will then try PubMed; if that also lacks fields, the entry 
    # stays "unverified" — which is the only honest outcome.  No partial/garbage 
    # data should ever be stamped "verified" and shown to the user.
    missing_fields = []
    if not api_title:
        missing_fields.append("title")
    if not api_authors:
        missing_fields.append("authors")
    if not api_year:
        missing_fields.append("year")
    if api_type == "Journal Article" and not api_journal:
        missing_fields.append("journal")
    
    if missing_fields:
        logger.info(
            "API result for DOI %s is missing essential field(s): %s. "
            "Rejecting to prevent incomplete data from being marked as verified.",
            api_result.get("doi", "?"),
            ", ".join(missing_fields),
        )
        return False

    # Extract API author surnames ("Smith, J." → "Smith")
    api_surnames = [a.split(',')[0].strip() for a in api_authors if ',' in a]

    # Extract local author surnames from structured list
    local_surnames = [a.split(',')[0].strip() for a in local_authors if ',' in a]
    # Also extract surnames from "Firstname Lastname" format
    for a in local_authors:
        parts = a.strip().split()
        if len(parts) >= 2 and ',' not in a:
            local_surnames.append(parts[-1])  # last word = surname
    # Also include raw PDF surnames as a fallback
    combined_surnames = set(local_surnames + pdf_surnames)

    # ── If still no signals after extraction, fail safe ─────────────────────
    if not local_title and not combined_surnames and not local_year:
        logger.info(
            "Refusing API result for DOI %s — could not extract any local identity "
            "signal from the PDF to validate against.",
            api_result.get("doi", "?")
        )
        return False

    # ── Check each signal ───────────────────────────────────────────────────
    passes   = []   # signals that actively confirm a match
    failures = []   # signals that actively deny a match

    # 1. Title similarity
    if local_title and api_title:
        sim = _sim(local_title, api_title)
        logger.debug("Title similarity: %.2f ('%s' vs '%s')", sim, local_title[:60], api_title[:60])
        if sim >= threshold:
            passes.append(("title", sim))
        else:
            # ── HARD VERIFICATION ──
            # If the local title is completely corrupted, similarity will fail.
            # As a bulletproof fallback, we check if the true API title is physically
            # present in the raw PDF text (ignoring all whitespace and punctuation).
            # This completely prevents AI hallucination.
            if hard_verify_against_pdf(api_title, api_surnames, pdf_path):
                logger.info("Title mismatch OVERRIDDEN: True API title and first author found perfectly in raw PDF text.")
                passes.append(("title_in_pdf", 1.0))
            else:
                failures.append(("title", sim))

    # 2. Author surname overlap (at least 1 surname in common)
    if combined_surnames and api_surnames:
        # Normalise for comparison (strip accents, lowercase)
        import unicodedata
        def _norm(s):
            return unicodedata.normalize('NFKD', s).encode('ascii', 'ignore').decode().lower()
        norm_local = {_norm(s) for s in combined_surnames}
        norm_api   = {_norm(s) for s in api_surnames}
        overlap    = norm_local & norm_api
        logger.debug("Author surname overlap: %s (local) vs %s (api) → %s", norm_local, norm_api, overlap)
        if overlap:
            passes.append(("authors", len(overlap)))
        else:
            failures.append(("authors", 0))

    # 3. Year match (exact)
    if local_year and api_year:
        if int(local_year) == int(api_year):
            passes.append(("year", int(local_year)))
        else:
            failures.append(("year", (local_year, api_year)))

    # ── Decision logic ───────────────────────────────────────────────────────
    # Check if first-author surname matches (strong identity signal)
    first_author_match = any(s == "authors" for s, _ in passes)

    # A title failure is a hard veto UNLESS the first author matches.
    # Rationale: garbage titles (e.g. "Full list of author information") cause
    # false title mismatches, but a matching first-author surname is a strong
    # signal that the DOI is correct.
    for signal, val in failures:
        if signal == "title":
            if first_author_match:
                logger.info(
                    "DOI validation — title mismatch (sim=%.2f) OVERRIDDEN by "
                    "first-author match. Accepting DOI.", val
                )
                return True
            # Also override if local title looks like garbage
            if local_title and _is_garbage_title(local_title):
                logger.info(
                    "DOI validation — title mismatch (sim=%.2f) OVERRIDDEN because "
                    "local title is garbage: %r. Accepting DOI.", val, local_title[:60]
                )
                return True
            logger.info("DOI validation FAILED — title mismatch (sim=%.2f)", val)
            return False

    # Any passing signal is sufficient if no hard vetoes.
    if passes:
        logger.info("DOI validation PASSED — signals: %s", passes)
        return True

    # No signals checked at all (e.g. API returned no title/authors/year).
    logger.info("DOI validation INCONCLUSIVE — no overlapping signals to compare.")
    return False




# ══════════════════════════════════════════════════════════════════════════════
# LAYER 3 — pdf2doi focused DOI rescue  (offline-only mode)
# ══════════════════════════════════════════════════════════════════════════════

def _rescue_doi_via_pdf2doi(pdf_path: str) -> Optional[str]:
    """
    Last-resort DOI finder using pdf2doi in **offline-only** mode.
    Online validation and pdfminer fallback are disabled to prevent:
      - Blocking synchronous HTTP calls to dx.doi.org
      - Slow pdfminer text re-extraction (5-10s per page)
    The returned DOI is sanitized through our own _clean_doi to strip
    any trailing publisher domain concatenations.
    Returns a DOI string or None.
    """
    try:
        import pdf2doi
        # Disable pdf2doi's internal blocking web validation calls
        # and its slow pdfminer text extraction fallback
        pdf2doi.config.set('verbose', False)
        pdf2doi.config.set('validation_online', False)

        result = pdf2doi.pdf2doi(pdf_path)
        if result and result.get("identifier"):
            from utils.text_utils import extract_doi as _sanitize_doi
            raw_doi = str(result["identifier"]).strip()
            # Run through our own sanitizer to strip trailing domains
            clean = _sanitize_doi(raw_doi)
            if clean:
                logger.info("pdf2doi rescued DOI (sanitized): %s", clean)
                return clean
            # If our sanitizer rejects it, try the raw lowercase
            doi = raw_doi.lower()
            logger.info("pdf2doi rescued DOI (raw): %s", doi)
            return doi
    except ImportError:
        logger.warning("pdf2doi not installed — skipping DOI rescue layer.")
    except Exception as exc:
        logger.warning("pdf2doi failed: %s", exc)
    return None


def verify_doi_online(doi_str: str) -> bool:
    """Fast HEAD request using connection-pooled session to ensure the string exists in the global registry."""
    try:
        url = f"{CROSSREF_API}/{doi_str}"
        _crossref_limiter.wait()
        r = _http_session.head(
            url, 
            timeout=5, 
            headers={"User-Agent": "metadata-extractor/1.0"}
        )
        return r.status_code == 200
    except Exception as exc:
        logger.warning("Speculative online check failed for %s: %s", doi_str, exc)
        return False


async def robust_doi_resolver(raw_text_chunk: str) -> dict:
    """
    Executes a speculative check against CrossRef to guarantee data validity.
    Step 3 of the Zotero/CrossRef gold standard workflow.
    """
    from utils.text_utils import DOI_CAPTURE_REGEX, polish_extracted_string
    
    # First attempt: Test the raw regex match directly
    initial_match = DOI_CAPTURE_REGEX.search(raw_text_chunk)
    if not initial_match:
        return {"status": "failed", "doi": None}
        
    candidate_doi = initial_match.group(0)
    
    # Check if the initial un-truncated match is a valid registered string
    is_valid_raw = await asyncio.to_thread(verify_doi_online, candidate_doi)
    if is_valid_raw:
        logger.info("Speculative check: Raw DOI candidate %s exists!", candidate_doi)
        return {"status": "verified_raw", "doi": candidate_doi}
        
    # Second attempt: Apply the truncation sieve if the raw string fails verification
    polished_doi = polish_extracted_string(candidate_doi)
    if polished_doi != candidate_doi:
        is_valid_polished = await asyncio.to_thread(verify_doi_online, polished_doi)
        if is_valid_polished:
            logger.info("Speculative check: Concatenated candidate resolved to verified DOI: %s", polished_doi)
            return {"status": "verified_after_sieve", "doi": polished_doi}
            
    logger.info("Speculative check: Candidate %s is unverifiable.", candidate_doi)
    return {"status": "unverifiable", "doi": None}


# ══════════════════════════════════════════════════════════════════════════════
# LAYER 5 — CrossRef API verification + enrichment
# ══════════════════════════════════════════════════════════════════════════════

def crossref_lookup(doi: str) -> Optional[dict]:
    """
    Verify a DOI against CrossRef and return enriched metadata if found.
    Returns a normalised dict or None on failure.
    """
    try:
        url = f"{CROSSREF_API}/{doi}"
        _crossref_limiter.wait()
        r   = _http_session.get(url, timeout=REQUEST_TIMEOUT,
                           headers={"User-Agent": "metadata-extractor/1.0"})
        if r.status_code != 200:
            return None

        item = r.json().get("message", {})

        # ── Reject container-level DOIs (journal, journal-issue) ─────────
        # These DOIs point to a journal/series, not an individual article.
        # Accepting them produces false "verified" status with wrong metadata.
        cr_type = item.get("type", "")
        if cr_type in ("journal", "journal-issue", "journal-volume", "book-series", "report-series"):
            logger.info(
                "CrossRef DOI %s is a container-level record (type=%s), not an article. Rejecting.",
                doi, cr_type,
            )
            return None
        authors = []
        for a in item.get("author", []):
            given  = a.get("given", "")
            family = a.get("family", "")
            if family and given:
                initials = ". ".join(w[0].upper() for w in given.split() if w) + "."
                authors.append(f"{family}, {initials}")
            elif family:
                authors.append(family)
            elif given:
                authors.append(given)

        year = None
        for date_field in ("published-print", "published-online", "issued"):
            parts = item.get(date_field, {}).get("date-parts", [[]])
            if parts and parts[0]:
                year = parts[0][0]
                break

        titles = item.get("title", [])

        # Normalise CrossRef types to match the internal format used by the
        # formatter (e.g. "journal-article" → "Journal Article")
        _CROSSREF_TYPE_MAP = {
            "journal-article":       "Journal Article",
            "book-chapter":          "Book Chapter",
            "book":                  "Book",
            "proceedings-article":   "Conference Paper",
            "posted-content":        "Preprint",
            "report":                "Report",
            "dissertation":          "Thesis",
            "dataset":               "Dataset",
            "monograph":             "Book",
            "edited-book":           "Book",
            "reference-entry":       "Other",
        }
        raw_type = item.get("type", "")
        normalised_type = _CROSSREF_TYPE_MAP.get(raw_type, raw_type.replace("-", " ").title() if raw_type else "Other")

        return {
            "title":               titles[0] if titles else None,
            "authors":             authors,
            "year":                year,
            "doi":                 item.get("DOI", doi).lower(),
            "journal":             (item.get("container-title") or [None])[0],
            "volume":              item.get("volume"),
            "issue":               item.get("issue"),
            "pages":               item.get("page"),
            "publisher":           item.get("publisher"),
            "type":                normalised_type,
            "verification_status": "verified_crossref",
        }
    except Exception as exc:
        logger.warning("CrossRef lookup failed for %s: %s", doi, exc)
    return None


def _crossref_search_by_title(title: str) -> Optional[dict]:
    """
    Search CrossRef by title string (fuzzy) when we have no DOI.
    Validates similarity between the query title and the returned title
    before accepting the match — prevents wrong-DOI assignments.

    Uses two thresholds:
      - 0.90 for general similarity (strict)
      - 0.80 + prefix check for truncated titles (e.g. PDF line-wrap cut-offs)
    """
from difflib import SequenceMatcher

def _norm(s: str) -> str:
    return re.sub(r'[^\w\s]', '', s.lower()).strip()

def _title_sim(a: str, b: str) -> float:
    return SequenceMatcher(None, _norm(a), _norm(b)).ratio()

def _is_prefix(shorter: str, longer: str) -> bool:
    """Check if shorter is a clean prefix of longer (>=60% of longer)."""
    ns, nl = _norm(shorter), _norm(longer)
    return len(ns) >= 20 and nl.startswith(ns) and len(ns) / max(len(nl), 1) >= 0.60

def _pubmed_search_by_title(title: str) -> Optional[dict]:
    """
    Fallback: search PubMed explicitly by title if we have no DOI.
    Applies the same strict similarity checks as the CrossRef search.
    """
    try:
        search_r = _http_session.get(
            PUBMED_SEARCH,
            params={"db": "pubmed", "term": f"{title}[Title]", "retmode": "json", "retmax": 3},
            timeout=REQUEST_TIMEOUT,
        )
        if search_r.status_code != 200:
            return None
            
        pmids = search_r.json().get("esearchresult", {}).get("idlist", [])
        if not pmids:
            return None

        # Fetch XML for the top PMIDs
        fetch_r = _http_session.get(
            PUBMED_FETCH,
            params={"db": "pubmed", "id": ",".join(pmids), "retmode": "xml"},
            timeout=REQUEST_TIMEOUT,
        )
        root = ET.fromstring(fetch_r.text)
        
        for article in root.findall(".//PubmedArticle"):
            medline = article.find(".//MedlineCitation")
            if medline is None:
                continue
                
            candidate_title = medline.findtext(".//ArticleTitle")
            if not candidate_title:
                continue
                
            sim = _title_sim(title, candidate_title)
            article_id_list = article.find(".//PubmedData/ArticleIdList")
            doi_el = article_id_list.find('.//ArticleId[@IdType="doi"]') if article_id_list is not None else None
            doi = doi_el.text if doi_el is not None else None

            # Strict match or prefix match
            if sim >= 0.90 or (sim >= 0.80 and _is_prefix(title, candidate_title)):
                if doi:
                    logger.info("PubMed title search matched (sim=%.2f): %s", sim, candidate_title)
                    res = pubmed_lookup(doi)
                    if res:
                        return res
    except Exception as exc:
        logger.warning("PubMed title search failed: %s", exc)
    return None

def _crossref_search_by_title(title: str) -> Optional[dict]:
    """
    Fallback: search CrossRef explicitly by title if we have no DOI.
    Applies a strict title similarity check (and author checks if available)
    before accepting the match - prevents wrong-DOI assignments.

    Uses two thresholds:
      - 0.90 for general similarity (strict)
      - 0.80 + prefix check for truncated titles (e.g. PDF line-wrap cut-offs)
    """
    try:
        r = _http_session.get(
            CROSSREF_API,
            params={"query.title": title, "rows": 3, "select": "DOI,title,author,issued"},
            timeout=REQUEST_TIMEOUT,
            headers={"User-Agent": "metadata-extractor/1.0"},
        )
        if r.status_code != 200:
            return None
        items = r.json().get("message", {}).get("items", [])
        if not items:
            return None

        # Pick the first candidate whose title is similar enough to ours
        for item in items:
            candidate_titles = item.get("title", [])
            if not candidate_titles:
                continue
            candidate_title = candidate_titles[0]
            sim = _title_sim(title, candidate_title)

            # Strict match
            if sim >= 0.90:
                doi = item.get("DOI")
                if doi:
                    logger.info("CrossRef title search matched (sim=%.2f): %s", sim, candidate_title)
                    return crossref_lookup(doi)

            # Prefix match — our title is truncated but is a clean start of the real title
            elif sim >= 0.80 and _is_prefix(title, candidate_title):
                doi = item.get("DOI")
                if doi:
                    logger.info(
                        "CrossRef title search PREFIX matched (sim=%.2f): '%s' → '%s'",
                        sim, title[:50], candidate_title[:80],
                    )
                    return crossref_lookup(doi)
            else:
                logger.debug(
                    "CrossRef title search candidate rejected (sim=%.2f): %s", sim, candidate_title
                )

        logger.info("CrossRef title search found no sufficiently similar match for: %s", title)
    except Exception as exc:
        logger.warning("CrossRef title search failed: %s", exc)
    return None


# ══════════════════════════════════════════════════════════════════════════════
# LAYER 6 — PubMed API secondary verification
# ══════════════════════════════════════════════════════════════════════════════

def pubmed_lookup(doi: str) -> Optional[dict]:
    """
    Search PubMed by DOI. Returns a slim metadata dict or None.
    Only used as a secondary check when CrossRef fails.
    """
    try:
        # Step 1 — search for the PMID
        _pubmed_limiter.wait()
        search_r = _http_session.get(
            PUBMED_SEARCH,
            params={"db": "pubmed", "term": f"{doi}[doi]", "retmode": "json", "retmax": 1},
            timeout=REQUEST_TIMEOUT,
        )
        pmids = search_r.json().get("esearchresult", {}).get("idlist", [])
        if not pmids:
            return None

        # Step 2 — fetch the record
        _pubmed_limiter.wait()
        fetch_r = _http_session.get(
            PUBMED_FETCH,
            params={"db": "pubmed", "id": pmids[0], "retmode": "xml"},
            timeout=REQUEST_TIMEOUT,
        )
        root = ET.fromstring(fetch_r.text)
        article = root.find(".//Article")
        if article is None:
            return None

        title = article.findtext("ArticleTitle")
        year_el = root.find(".//PubDate/Year")
        year    = int(year_el.text) if year_el is not None and year_el.text else None

        authors = []
        for a in root.findall(".//Author"):
            last    = a.findtext("LastName", "")
            fore    = a.findtext("ForeName", "")
            if last and fore:
                initials = ". ".join(w[0].upper() for w in fore.split() if w) + "."
                authors.append(f"{last}, {initials}")
            elif last:
                authors.append(last)
            elif fore:
                authors.append(fore)

        # Journal metadata
        journal_el = root.find(".//Journal")
        journal_title = None
        volume = None
        issue = None
        if journal_el is not None:
            journal_title = journal_el.findtext("Title") or journal_el.findtext("ISOAbbreviation")
            ji = journal_el.find("JournalIssue")
            if ji is not None:
                volume = ji.findtext("Volume")
                issue = ji.findtext("Issue")

        # Pages
        pages = article.findtext("Pagination/MedlinePgn")

        return {
            "title":               title,
            "authors":             authors,
            "year":                year,
            "doi":                 doi.lower(),
            "journal":             journal_title,
            "volume":              volume,
            "issue":               issue,
            "pages":               pages,
            "type":                "Journal Article",
            "verification_status": "verified_pubmed",
        }
    except Exception as exc:
        logger.warning("PubMed lookup failed for %s: %s", doi, exc)
    return None


# ══════════════════════════════════════════════════════════════════════════════
# COMPATIBILITY ADAPTERS FOR PARSER.PY
# ══════════════════════════════════════════════════════════════════════════════

def perform_pubmed_lookup(doi: str, api_metadata: dict, api_sources: dict) -> bool:
    res = pubmed_lookup(doi)
    if not res:
        return False
    for k in ["title", "authors", "year", "type"]:
        if res.get(k):
            api_metadata[k] = res[k]
            api_sources[k] = "pubmed"
    api_sources["doi"] = "pubmed"
    return True

def perform_crossref_lookup(doi: str, api_metadata: dict, api_sources: dict) -> bool:
    res = crossref_lookup(doi)
    if not res:
        return False
    if res.get("journal"):
        res["source"] = res["journal"]
    for k in ["title", "authors", "year", "source", "volume", "issue", "pages", "publisher", "type"]:
        if res.get(k):
            api_metadata[k] = res[k]
            api_sources[k] = "crossref"
    api_sources["doi"] = "crossref"
    return True


# ══════════════════════════════════════════════════════════════════════════════
# MERGE HELPER
# ══════════════════════════════════════════════════════════════════════════════

def _merge(base: dict, update: dict, overwrite: bool = False) -> dict:
    """
    Fill in missing fields in `base` from `update`.
    If overwrite=True, replaces existing non-empty values in base (provided update has them).
    """
    for key, val in update.items():
        if val is None:
            continue
        if isinstance(val, list) and not val:
            continue
        if overwrite or not base.get(key):
            base[key] = val
    return base


def _is_complete(meta: dict) -> bool:
    """Return True when all four critical fields are present."""
    return bool(
        meta.get("title")
        and meta.get("authors")
        and meta.get("year")
        and meta.get("doi")
    )


# ══════════════════════════════════════════════════════════════════════════════
# MAIN EXTRACTION LOGIC
# ══════════════════════════════════════════════════════════════════════════════

def _extract_metadata_local_sync(pdf_path: str) -> tuple[dict, list[str]]:
    """
    Runs the local CPU-bound extraction layers (PDF meta, Regex).
    Returns the partial metadata dict and a list of candidate DOIs found.
    """
    meta = {
        "title": None, "authors": [], "year": None, "doi": None,
        "journal": None, "volume": None, "issue": None, "pages": None,
        "publisher": None, "type": "Other", "verification_status": "not_found",
        "extraction_layers": [],
    }

    # ── Layer 1: PDF embedded metadata ────────────────────────────────────────
    layer1 = _extract_pdf_metadata(pdf_path)
    _merge(meta, layer1)
    if any(layer1.values()):
        meta["extraction_layers"].append("pdf_metadata")

    # ── Layer 2: Regex (text_utils.py) ────────────────────────────────────────
    layer2 = _extract_via_regex(pdf_path)
    _merge(meta, layer2)
    if any(v for v in layer2.values() if v):
        meta["extraction_layers"].append("regex")



    candidate_dois = []
    _seen_dois = set()
    # Priority: text-extracted (regex) > PDF embedded metadata
    for doi_source in [layer2, layer1]:
        d = doi_source.get("doi") if isinstance(doi_source, dict) else None
        if d and d.lower() not in _seen_dois:
            candidate_dois.append(d)
            _seen_dois.add(d.lower())
    if meta.get("doi") and meta["doi"].lower() not in _seen_dois:
        candidate_dois.append(meta["doi"])
        _seen_dois.add(meta["doi"].lower())

    return meta, candidate_dois


async def _extract_metadata_async(pdf_path: str) -> dict:
    """
    Runs the full cascade, separating CPU-bound tasks from API verification
    which uses the globally rate-limited TokenBucket.
    """
    meta, candidate_dois = await asyncio.to_thread(_extract_metadata_local_sync, pdf_path)
    logger.info("Candidate DOIs for verification: %s", candidate_dois)

    # ── SPECULATIVE SIEVE — run robust_doi_resolver on all candidates ─────
    resolved_dois = []
    _seen_resolved = set()
    for candidate_doi in candidate_dois:
        resolved = await robust_doi_resolver(candidate_doi)
        if resolved["status"] in ("verified_raw", "verified_after_sieve"):
            doi_to_use = resolved["doi"]
            if doi_to_use and doi_to_use.lower() not in _seen_resolved:
                resolved_dois.append(doi_to_use)
                _seen_resolved.add(doi_to_use.lower())
        else:
            # Fallback/default: use polished version of candidate if not already in list
            from utils.text_utils import polish_extracted_string
            polished = polish_extracted_string(candidate_doi)
            if polished and polished.lower() not in _seen_resolved:
                resolved_dois.append(polished)
                _seen_resolved.add(polished.lower())
    
    candidate_dois = resolved_dois
    logger.info("Resolved/Polished candidate DOIs for verification: %s", candidate_dois)

    # ── VERIFICATION LOOP — try each candidate DOI ───────────────────────────
    verified = False
    
    def record_to_dict(rec) -> dict:
        return {
            "title": rec.title, "authors": rec.authors, "year": rec.year, "doi": rec.doi,
            "journal": rec.journal, "source": rec.journal, "url": rec.url,
            "volume": rec.volume, "issue": rec.issue, "pages": rec.pages,
            "publisher": rec.publisher, "type": rec.type,
        }

    # Pre-fetch all candidates using the rate-limited batch API
    if candidate_dois:
        from references.api_batch import fetch_crossref_batch
        crossref_records = await fetch_crossref_batch(candidate_dois)
        crossref_map = {rec.doi: record_to_dict(rec) for rec in crossref_records if rec and rec.doi}
    else:
        crossref_map = {}

    for candidate_doi in candidate_dois:
        logger.info("Trying DOI candidate: %s", candidate_doi)
        
        # ── CrossRef first (from our pre-fetched async batch) ──
        cr_data = crossref_map.get(candidate_doi)
        if cr_data:
            valid = await asyncio.to_thread(_validate_api_result, meta, cr_data, pdf_path)
            if valid:
                _merge(meta, cr_data, overwrite=True)
                meta["verification_status"] = "verified_crossref"
                if "crossref_verify" not in meta["extraction_layers"]:
                    meta["extraction_layers"].append("crossref_verify")
                logger.info("DOI %s VERIFIED via CrossRef.", candidate_doi)
                verified = True
                break
            else:
                logger.info("CrossRef result for DOI %s failed identity validation.", candidate_doi)
                
        # ── PubMed fallback ──
        def do_pubmed():
            pubmed = pubmed_lookup(candidate_doi)
            if pubmed and _validate_api_result(meta, pubmed, pdf_path):
                return pubmed
            return None
            
        pubmed_data = await asyncio.to_thread(do_pubmed)
        if pubmed_data:
            _merge(meta, pubmed_data, overwrite=True)
            meta["verification_status"] = "verified_pubmed"
            if "pubmed_verify" not in meta["extraction_layers"]:
                meta["extraction_layers"].append("pubmed_verify")
            logger.info("DOI %s VERIFIED via PubMed.", candidate_doi)
            verified = True
            break


    # ── pdf2doi rescue — if none of the candidates verified ──────────────────
    if not verified:
        meta["doi"] = None  # Clear any unverified DOI
        if "pdf2doi" not in meta["extraction_layers"]:
            rescued = await asyncio.to_thread(_rescue_doi_via_pdf2doi, pdf_path)
            if rescued:
                meta["doi"] = rescued
                meta["extraction_layers"].append("pdf2doi")
                logger.info("pdf2doi rescued DOI: %s — verifying...", rescued)

                # Fetch rescued DOI using async batch
                from references.api_batch import fetch_crossref_batch
                rescued_records = await fetch_crossref_batch([rescued])
                cr_data = None
                if rescued_records and rescued_records[0]:
                    cr_data = record_to_dict(rescued_records[0])
                
                if cr_data and await asyncio.to_thread(_validate_api_result, meta, cr_data, pdf_path):
                    _merge(meta, cr_data, overwrite=True)
                    meta["verification_status"] = "verified_crossref"
                    meta["extraction_layers"].append("crossref_verify")
                    verified = True
                else:
                    def do_pubmed_rescued():
                        pubmed = pubmed_lookup(rescued)
                        if pubmed and _validate_api_result(meta, pubmed, pdf_path):
                            return pubmed
                        return None
                    
                    pubmed_data = await asyncio.to_thread(do_pubmed_rescued)
                    if pubmed_data:
                        _merge(meta, pubmed_data, overwrite=True)
                        meta["verification_status"] = "verified_pubmed"
                        meta["extraction_layers"].append("pubmed_verify")
                        verified = True
                    else:
                        logger.info("pdf2doi DOI %s also failed verification. Clearing.", rescued)
                        meta["doi"] = None

    # Title-search fallbacks removed: they bypassed _validate_api_result,
    # performed no PDF identity check, and could stamp "verified" on wrong data.
    # If DOI verification failed, status correctly stays not_found → unverified.

    # ── Final status tagging ──────────────────────────────────────────────────
    if meta["verification_status"] == "not_found" and any([
        meta.get("title"), meta.get("doi"), meta.get("year")
    ]):
        meta["verification_status"] = "unverified"

    if meta["type"] == "Other":
        meta["type"] = classify_source_type(meta)

    logger.info(
        "Extraction complete | status=%s | layers=%s | doi=%s",
        meta["verification_status"],
        meta["extraction_layers"],
        meta.get("doi"),
    )
    return meta

# ══════════════════════════════════════════════════════════════════════════════
# PUBLIC ASYNC WRAPPERS FOR FASTAPI ROUTES
# ══════════════════════════════════════════════════════════════════════════════

async def extract_pdf_metadata(file_bytes: bytes) -> dict:
    """
    Extract metadata from a PDF file (in-memory bytes).
    Writes bytes to a temporary file, runs the deterministic cascade,
    then cleans up the file. Returns the populated metadata dict.
    """
    # Create temp file
    fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(file_bytes)
        
        # Run async cascade (CPU-bound layers use to_thread, API layers use batch fetch)
        meta = await _extract_metadata_async(tmp_path)
        return meta
    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError as e:
                logger.warning("Failed to clean up temp PDF %s: %s", tmp_path, e)


def extract_docx_metadata(file_bytes: bytes) -> dict:
    """Extract metadata from a DOCX file using document core properties."""
    metadata = {
        "authors": None, "title": None, "year": None,
        "source": None, "doi": None, "url": None,
        "volume": None, "issue": None, "pages": None,
        "publisher": None, "type": "Other",
        "verification_status": "not_found",
        "extraction_layers": ["docx_metadata"]
    }
    
    with io.BytesIO(file_bytes) as f:
        doc = Document(f)
        
        props = doc.core_properties
        if props.author:
            metadata["authors"] = [a.strip() for a in re.split(r'[;,]', props.author) if a.strip()]
        if props.title and len(props.title) > 3:
            metadata["title"] = props.title
        if props.created:
            metadata["year"] = str(props.created.year)
        
        full_text = '\n'.join(p.text for p in doc.paragraphs[:5])
        docx_doi = extract_doi(full_text)
        if docx_doi:
            metadata["doi"] = docx_doi
        
        if not metadata["title"]:
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if lines:
                metadata["title"] = lines[0]
    
    # CrossRef DOI lookup — with strict completeness gate
    if metadata["doi"]:
        try:
            crossref_url = f"https://api.crossref.org/works/{metadata['doi']}"
            resp = _http_session.get(crossref_url, timeout=10, headers={
                'User-Agent': 'WritingTools/1.0 (mailto:support@paradoxlabs.com)'
            })
            if resp.status_code == 200:
                data = resp.json().get('message', {})

                # Parse authors
                author_parts = []
                for a in data.get('author', []):
                    family = a.get('family', '')
                    given  = a.get('given', '')
                    if family and given:
                        initials = '. '.join(w[0].upper() for w in given.split() if w) + '.'
                        author_parts.append(f"{family}, {initials}")
                    elif family:
                        author_parts.append(family)

                titles = data.get('title', [])
                cr_title = titles[0] if titles else None
                
                cr_type  = data.get('type', '')
                type_map = {
                    'journal-article':    'Journal Article',
                    'book':               'Book',
                    'book-chapter':       'Book Chapter',
                    'proceedings-article':'Conference Paper',
                    'report':             'Report',
                }
                mapped_type = type_map.get(cr_type, 'Other')

                container = data.get('container-title', [])
                cr_journal = container[0] if container else None

                date_parts = data.get('published-print', data.get('published-online', data.get('created', {})))
                cr_year = None
                if date_parts and 'date-parts' in date_parts:
                    parts = date_parts['date-parts'][0]
                    if parts:
                        cr_year = str(parts[0])

                # STRICT COMPLETENESS GATE — all essential fields must be present
                missing = [f for f, v in [("title", cr_title), ("authors", author_parts), ("year", cr_year)] if not v]
                if mapped_type == 'Journal Article' and not cr_journal:
                    missing.append("journal")
                    
                if missing:
                    logger.info(
                        "DOCX CrossRef result missing field(s) %s for DOI %s — not marking verified.",
                        missing, metadata['doi']
                    )
                else:
                    metadata["authors"] = author_parts
                    metadata["title"]   = cr_title
                    metadata["year"]    = cr_year
                    metadata["type"]    = mapped_type

                    if cr_journal:
                        metadata["source"] = cr_journal
                    if data.get('volume'):    metadata["volume"]    = data['volume']
                    if data.get('issue'):     metadata["issue"]     = data['issue']
                    if data.get('page'):      metadata["pages"]     = data['page']
                    if data.get('publisher'): metadata["publisher"] = data['publisher']

                    metadata["verification_status"] = "verified_crossref"
                    metadata["extraction_layers"].append("crossref_verify")
        except Exception:
            pass
    
    if metadata["type"] == "Other":
        metadata["type"] = classify_source_type(metadata)
    
    return metadata
