"""
Document conversion API routes and helper tools.
"""
import io
import os
import shutil
import asyncio
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import List

from fastapi import APIRouter, UploadFile, File, HTTPException
from starlette.responses import StreamingResponse
from PyPDF2 import PdfReader, PdfWriter
from docx import Document

from core.config import (
    PDF2DOCX_AVAILABLE, Pdf2DocxConverter,
    PILLOW_AVAILABLE, Image,
    PDF2IMAGE_AVAILABLE, convert_from_bytes,
    TESSERACT_AVAILABLE, pytesseract,
    POPPLER_PATH,
)
from utils.text_extraction import extract_pdf_text

router = APIRouter()


# ─── Helper tools ─────────────────────────────────────────────────────────

def _find_libreoffice():
    """Find the LibreOffice soffice executable."""
    candidates = [
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        shutil.which('soffice'),
    ]
    for c in candidates:
        if c and os.path.isfile(c):
            return c
    return None


def _ocr_pdf_to_text(pdf_bytes: bytes) -> str:
    """Use pytesseract to OCR a scanned PDF."""
    if not TESSERACT_AVAILABLE or not PDF2IMAGE_AVAILABLE:
        return ""
    try:
        images = convert_from_bytes(pdf_bytes, dpi=300, poppler_path=POPPLER_PATH)
        text_parts = []
        for img in images:
            page_text = pytesseract.image_to_string(img)
            text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"   [!] OCR failed: {e}")
        return ""


def _ocr_pdf_to_docx(pdf_bytes: bytes, docx_path: str):
    """
    OCR a scanned PDF and create a DOCX that preserves table structure.
    Uses PyMuPDF's built-in OCR.
    """
    import fitz  # PyMuPDF
    from docx.shared import Pt

    env_path = os.environ.get('PATH', '')
    tess_dir = r'C:\Program Files\Tesseract-OCR'
    if os.path.isdir(tess_dir) and tess_dir not in env_path:
        os.environ['PATH'] = tess_dir + ';' + env_path

    try:
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        print(f"   [i] Opened PDF: {len(pdf_doc)} pages")

        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(10)
        style.font.name = 'Calibri'

        for page_idx in range(len(pdf_doc)):
            page = pdf_doc[page_idx]
            if page_idx > 0:
                doc.add_page_break()

            print(f"   [i] OCR'ing page {page_idx + 1}/{len(pdf_doc)}...")

            tp = page.get_textpage_ocr(language="eng", dpi=300, full=True)
            words = page.get_text("words", textpage=tp)
            print(f"   [i]   OCR found {len(words)} words")

            if not words:
                continue

            page_width = page.rect.width

            # Group words into visual lines
            line_tolerance = 5
            lines = []
            current_line = [words[0]]

            for w in words[1:]:
                avg_y = sum(ww[1] for ww in current_line) / len(current_line)
                if abs(w[1] - avg_y) <= line_tolerance:
                    current_line.append(w)
                else:
                    current_line.sort(key=lambda x: x[0])
                    lines.append(current_line)
                    current_line = [w]
            if current_line:
                current_line.sort(key=lambda x: x[0])
                lines.append(current_line)

            # Detect column structure
            min_gap = page_width * 0.03
            all_gap_midpoints = []

            for line in lines:
                if len(line) < 2:
                    continue
                for i in range(len(line) - 1):
                    gap = line[i + 1][0] - line[i][2]
                    if gap > min_gap:
                        midpoint = (line[i][2] + line[i + 1][0]) / 2
                        all_gap_midpoints.append(midpoint)

            strong_boundaries = []
            if all_gap_midpoints:
                all_gap_midpoints.sort()
                cluster_tolerance = page_width * 0.04
                clusters = []
                current_cluster = [all_gap_midpoints[0]]

                for gm in all_gap_midpoints[1:]:
                    if gm - current_cluster[-1] <= cluster_tolerance:
                        current_cluster.append(gm)
                    else:
                        clusters.append(current_cluster)
                        current_cluster = [gm]
                clusters.append(current_cluster)

                min_appearances = min(3, max(2, len(lines) // 5))
                for cluster in clusters:
                    if len(cluster) >= min_appearances:
                        strong_boundaries.append(sorted(cluster)[len(cluster) // 2])

            if strong_boundaries:
                num_cols = len(strong_boundaries) + 1
                col_ranges = []
                prev = 0
                for b in strong_boundaries:
                    col_ranges.append((prev, b))
                    prev = b
                col_ranges.append((prev, page_width + 50))

                def assign_to_columns(line_words):
                    row = [''] * num_cols
                    for w in line_words:
                        word_center = (w[0] + w[2]) / 2
                        for ci, (cl, cr) in enumerate(col_ranges):
                            if cl <= word_center < cr:
                                row[ci] = (row[ci] + ' ' + w[4]).strip() if row[ci] else w[4]
                                break
                        else:
                            row[-1] = (row[-1] + ' ' + w[4]).strip() if row[-1] else w[4]
                    return row

                min_cols_for_table = min(3, num_cols)
                blocks = []
                pending_text = []
                pending_table = []

                for line in lines:
                    row = assign_to_columns(line)
                    populated = sum(1 for c in row if c.strip())

                    if populated >= min_cols_for_table:
                        if pending_text:
                            blocks.append(('text', pending_text))
                            pending_text = []
                        pending_table.append(row)
                    else:
                        if pending_table:
                            blocks.append(('table', pending_table))
                            pending_table = []
                        line_text = ' '.join(w[4] for w in line)
                        if line_text.strip():
                            pending_text.append(line_text.strip())

                if pending_table:
                    blocks.append(('table', pending_table))
                if pending_text:
                    blocks.append(('text', pending_text))

                for btype, bdata in blocks:
                    if btype == 'text':
                        for text_line in bdata:
                            if text_line:
                                p = doc.add_paragraph(text_line)
                                if text_line.isupper() and len(text_line) > 3:
                                    for run in p.runs:
                                        run.bold = True
                    elif btype == 'table':
                        if not bdata:
                            continue
                        max_col = 0
                        for row in bdata:
                            for ci in range(len(row) - 1, -1, -1):
                                if row[ci].strip():
                                    max_col = max(max_col, ci + 1)
                                    break
                        if max_col < 2:
                            for row in bdata:
                                text = ' '.join(c for c in row if c.strip())
                                if text:
                                    doc.add_paragraph(text)
                            continue

                        table = doc.add_table(rows=len(bdata), cols=max_col)
                        table.style = 'Table Grid'
                        for r_idx, row_data in enumerate(bdata):
                            for c_idx in range(max_col):
                                cell = table.cell(r_idx, c_idx)
                                cell.text = (row_data[c_idx] if c_idx < len(row_data) else '').strip()
                                if r_idx == 0:
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            run.bold = True
                        doc.add_paragraph('')
            else:
                for line in lines:
                    line_text = ' '.join(w[4] for w in line)
                    if line_text.strip():
                        p = doc.add_paragraph(line_text.strip())
                        if line_text.strip().isupper() and len(line_text.strip()) > 3:
                            for run in p.runs:
                                run.bold = True

        pdf_doc.close()
        doc.save(docx_path)

        final_doc = Document(docx_path)
        total_text = "\n".join(p.text for p in final_doc.paragraphs).strip()
        total_tables = len(final_doc.tables)
        print(f"   [✓] OCR complete: {len(total_text)} chars, {total_tables} tables")

        return True

    except Exception as e:
        print(f"   [!] PyMuPDF OCR failed: {e}")
        import traceback
        traceback.print_exc()

        if TESSERACT_AVAILABLE and PDF2IMAGE_AVAILABLE:
            try:
                images = convert_from_bytes(pdf_bytes, dpi=300, poppler_path=POPPLER_PATH)
                doc = Document()
                for page_idx, img in enumerate(images):
                    if page_idx > 0:
                        doc.add_page_break()
                    page_text = pytesseract.image_to_string(img)
                    for para in page_text.split('\n'):
                        if para.strip():
                            doc.add_paragraph(para.strip())
                doc.save(docx_path)
                return True
            except:
                pass
        return False


# ─── Endpoints ────────────────────────────────────────────────────────────

@router.get("/api/health")
async def health_check():
    return {"status": "ok"}


@router.get("/api/convert/capabilities")
async def conversion_capabilities():
    """Return which conversion tools are available on this server."""
    libre = _find_libreoffice()
    return {
        "pdf_to_word": PDF2DOCX_AVAILABLE,
        "word_to_pdf": libre is not None,
        "pdf_to_text": True,
        "image_to_pdf": PILLOW_AVAILABLE,
        "pdf_to_images": PDF2IMAGE_AVAILABLE,
        "merge_pdf": True,
        "compress_pdf": True,
        "ocr": TESSERACT_AVAILABLE and PDF2IMAGE_AVAILABLE,
        "libreoffice_path": libre,
    }


@router.post("/api/convert/pdf-to-word")
async def convert_pdf_to_word(file: UploadFile = File(...)):
    """Convert PDF to DOCX."""
    if not PDF2DOCX_AVAILABLE:
        raise HTTPException(status_code=503, detail="pdf2docx is not installed on this server.")

    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Please upload a PDF file.")

    file_bytes = await file.read()
    if file_bytes[:4] != b'%PDF':
        raise HTTPException(status_code=400, detail="Not a valid PDF file.")

    tmp_dir = tempfile.mkdtemp()
    try:
        pdf_path = os.path.join(tmp_dir, "input.pdf")
        docx_path = os.path.join(tmp_dir, "output.docx")

        with open(pdf_path, 'wb') as f:
            f.write(file_bytes)

        cv = Pdf2DocxConverter(pdf_path)
        cv.convert(docx_path)
        cv.close()

        doc = Document(docx_path)
        text_content = "\n".join(p.text for p in doc.paragraphs).strip()

        if len(text_content) < 50 and TESSERACT_AVAILABLE and PDF2IMAGE_AVAILABLE:
            _ocr_pdf_to_docx(file_bytes, docx_path)

        with open(docx_path, 'rb') as f:
            result_bytes = f.read()

        out_name = Path(file.filename).stem + ".docx"
        return StreamingResponse(
            io.BytesIO(result_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{out_name}"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


@router.post("/api/convert/word-to-pdf")
async def convert_word_to_pdf(file: UploadFile = File(...)):
    """Convert DOCX to PDF using LibreOffice headless."""
    soffice = _find_libreoffice()
    if not soffice:
        raise HTTPException(status_code=503, detail="LibreOffice is not installed.")

    ext = Path(file.filename).suffix.lower()
    if ext not in ('.docx', '.doc', '.odt', '.rtf'):
        raise HTTPException(status_code=400, detail="Please upload a Word document (.docx, .doc, .odt, or .rtf).")

    file_bytes = await file.read()
    tmp_dir = tempfile.mkdtemp()
    try:
        in_path = os.path.join(tmp_dir, f"input{ext}")
        with open(in_path, 'wb') as f:
            f.write(file_bytes)

        result = await asyncio.to_thread(
            subprocess.run,
            [soffice, '--headless', '--convert-to', 'pdf', '--outdir', tmp_dir, in_path],
            capture_output=True, text=True, timeout=120
        )

        if result.returncode != 0:
            raise HTTPException(status_code=500, detail=f"LibreOffice conversion failed: {result.stderr}")

        pdf_path = os.path.join(tmp_dir, "input.pdf")
        if not os.path.isfile(pdf_path):
            pdf_files = [f for f in os.listdir(tmp_dir) if f.endswith('.pdf')]
            if pdf_files:
                pdf_path = os.path.join(tmp_dir, pdf_files[0])
            else:
                raise HTTPException(status_code=500, detail="LibreOffice did not produce a PDF output.")

        with open(pdf_path, 'rb') as f:
            result_bytes = f.read()

        out_name = Path(file.filename).stem + ".pdf"
        return StreamingResponse(
            io.BytesIO(result_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{out_name}"'}
        )
    except HTTPException:
        raise
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=504, detail="Conversion timed out.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


@router.post("/api/convert/pdf-to-text")
async def convert_pdf_to_text(file: UploadFile = File(...)):
    """Extract text from PDF. Falls back to OCR for scanned PDFs."""
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Please upload a PDF file.")

    file_bytes = await file.read()
    if file_bytes[:4] != b'%PDF':
        raise HTTPException(status_code=400, detail="Not a valid PDF file.")

    text = extract_pdf_text(file_bytes)

    if len(text.strip()) < 50 and TESSERACT_AVAILABLE and PDF2IMAGE_AVAILABLE:
        ocr_text = _ocr_pdf_to_text(file_bytes)
        if ocr_text.strip():
            text = ocr_text

    out_name = Path(file.filename).stem + ".txt"
    return StreamingResponse(
        io.BytesIO(text.encode('utf-8')),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'}
    )


@router.post("/api/convert/image-to-pdf")
async def convert_image_to_pdf(files: List[UploadFile] = File(...)):
    """Convert one or more images to a single PDF."""
    if not PILLOW_AVAILABLE:
        raise HTTPException(status_code=503, detail="Pillow/reportlab not installed.")

    images = []
    for f in files:
        ext = Path(f.filename).suffix.lower()
        if ext not in ('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp'):
            raise HTTPException(status_code=400, detail=f"Unsupported image format: {f.filename}.")
        img_bytes = await f.read()
        try:
            img = Image.open(io.BytesIO(img_bytes))
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            images.append(img)
        except Exception:
            raise HTTPException(status_code=400, detail=f"Could not read image: {f.filename}")

    if not images:
        raise HTTPException(status_code=400, detail="No valid images provided.")

    buf = io.BytesIO()
    if len(images) == 1:
        images[0].save(buf, 'PDF', resolution=150.0)
    else:
        images[0].save(buf, 'PDF', resolution=150.0, save_all=True, append_images=images[1:])
    buf.seek(0)

    out_name = Path(files[0].filename).stem + ".pdf" if len(files) == 1 else "images_combined.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'}
    )


@router.post("/api/convert/pdf-to-images")
async def convert_pdf_to_images(file: UploadFile = File(...)):
    """Convert each PDF page to a JPG image. Returns a ZIP archive."""
    if not PDF2IMAGE_AVAILABLE:
        raise HTTPException(status_code=503, detail="pdf2image is not installed.")

    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Please upload a PDF file.")

    file_bytes = await file.read()
    if file_bytes[:4] != b'%PDF':
        raise HTTPException(status_code=400, detail="Not a valid PDF file.")

    try:
        images = await asyncio.to_thread(
            convert_from_bytes, file_bytes, dpi=200, fmt='jpeg', poppler_path=POPPLER_PATH
        )
    except Exception as e:
        error_msg = str(e)
        if 'poppler' in error_msg.lower() or 'pdftoppm' in error_msg.lower():
            raise HTTPException(status_code=503, detail="Poppler is not installed.")
        raise HTTPException(status_code=500, detail=f"PDF to images failed: {error_msg}")

    zip_buf = io.BytesIO()
    stem = Path(file.filename).stem
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for i, img in enumerate(images, 1):
            img_buf = io.BytesIO()
            img.save(img_buf, 'JPEG', quality=90)
            img_buf.seek(0)
            zf.writestr(f"{stem}_page_{i}.jpg", img_buf.read())
    zip_buf.seek(0)

    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{stem}_images.zip"'}
    )


@router.post("/api/convert/merge-pdf")
async def merge_pdfs(files: List[UploadFile] = File(...)):
    """Merge multiple PDF files into one."""
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Please upload at least 2 PDF files to merge.")

    writer = PdfWriter()
    for f in files:
        if not f.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail=f"All files must be PDFs. '{f.filename}' is not a PDF.")
        file_bytes = await f.read()
        if file_bytes[:4] != b'%PDF':
            raise HTTPException(status_code=400, detail=f"'{f.filename}' is not a valid PDF file.")
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                writer.add_page(page)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading '{f.filename}': {str(e)}")

    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="merged.pdf"'}
    )


@router.post("/api/convert/compress-pdf")
async def compress_pdf(file: UploadFile = File(...)):
    """Compress a PDF by removing metadata and re-encoding streams."""
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Please upload a PDF file.")

    file_bytes = await file.read()
    if file_bytes[:4] != b'%PDF':
        raise HTTPException(status_code=400, detail="Not a valid PDF file.")

    original_size = len(file_bytes)

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        writer = PdfWriter()

        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)

        writer.add_metadata({
            '/Producer': 'WritingTools',
            '/Creator': 'WritingTools PDF Compressor',
        })

        buf = io.BytesIO()
        writer.write(buf)
        compressed_size = buf.tell()
        buf.seek(0)

        out_name = Path(file.filename).stem + "_compressed.pdf"
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{out_name}"',
                "X-Original-Size": str(original_size),
                "X-Compressed-Size": str(compressed_size),
                "X-Compression-Ratio": f"{(1 - compressed_size / original_size) * 100:.1f}%" if original_size > 0 else "0%",
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Compression failed: {str(e)}")
